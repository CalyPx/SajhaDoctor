from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

doc = Document()

# ─── Page Setup (A4, margins as per guidelines) ───
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Inches(1.5)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

# Helper function
def add_paragraph(text, size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, 
                  space_before=0, space_after=0, font_name='Times New Roman', color=None):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return p

# ═══════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════

# Top spacing
add_paragraph('', size=12, space_after=60)

# Title
add_paragraph('Earthquake Building Damage Prediction', size=16, bold=True, space_after=0)
add_paragraph('Using Machine Learning Algorithms', size=16, bold=True, space_after=24)

# Subtitle
add_paragraph('A Project Proposal', size=14, bold=True, space_before=30, space_after=12)

# Fulfillment line
add_paragraph('Submitted in partial fulfillment of the requirements', size=12, space_before=20, space_after=0)
add_paragraph('for the completion of the', size=12, space_after=0)
add_paragraph('Artificial Intelligence & Machine Learning Microdegree', size=12, bold=True, space_after=40)

# Submitted by
add_paragraph('Submitted by:', size=12, bold=True, space_before=40, space_after=8)
add_paragraph('Rohit Poudel', size=12, space_after=2)
add_paragraph('Enrollment No.: _______________', size=11, space_after=12)
add_paragraph('Keshav KC', size=12, space_after=2)
add_paragraph('Enrollment No.: _______________', size=11, space_after=30)

# Supervised by
add_paragraph('Supervised by:', size=12, bold=True, space_before=20, space_after=8)
add_paragraph('Mr. Diwash Sapkota', size=12, space_after=30)

# Institution
add_paragraph('Code for Change, Pokhara', size=13, bold=True, space_before=30, space_after=40)

# Date
add_paragraph('Submission Date: 23/04/2026', size=12, space_before=20)

# ─── Page break after cover ───
doc.add_page_break()

# ═══════════════════════════════════════════
# ACKNOWLEDGEMENT
# ═══════════════════════════════════════════

add_paragraph('ACKNOWLEDGEMENT', size=16, bold=True, space_after=24)

ack_text = (
    'We would like to express our sincere gratitude to our supervisor, '
    'Mr. Diwash Sapkota, for his continuous guidance and support throughout '
    'the development of this project proposal. His valuable feedback and '
    'encouragement helped us shape our ideas into a structured plan.'
)
add_paragraph(ack_text, size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)

ack_text2 = (
    'We are also thankful to Code for Change, Pokhara for providing us with '
    'the opportunity to be a part of the AI/ML Microdegree Program and for '
    'creating a learning environment where we could explore the practical '
    'applications of machine learning.'
)
add_paragraph(ack_text2, size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)

ack_text3 = (
    'We extend our appreciation to the open-source community and platforms '
    'like DrivenData and Kaggle for making real-world datasets publicly '
    'accessible, which made this project possible. Lastly, we would like '
    'to thank our families and friends for their constant motivation.'
)
add_paragraph(ack_text3, size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=24)

add_paragraph('Rohit Poudel', size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=30, space_after=2)
add_paragraph('Keshav KC', size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
add_paragraph('April 2026', size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

doc.add_page_break()

# ═══════════════════════════════════════════
# STUDENT'S DECLARATION
# ═══════════════════════════════════════════

add_paragraph("STUDENT'S DECLARATION", size=16, bold=True, space_after=24)

decl_text = (
    'We hereby declare that we are the sole authors of this project proposal '
    'and that no sources other than those referenced herein have been used. '
    'The work presented is original and any resemblance to existing projects '
    'is purely coincidental. We acknowledge that academic dishonesty will '
    'result in disqualification.'
)
add_paragraph(decl_text, size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=40)

# Signature lines
for name in ['Rohit Poudel', 'Keshav KC']:
    add_paragraph('___________________________', size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=30, space_after=2)
    add_paragraph(name, size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_paragraph('Enrollment No.: _______________', size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_paragraph('Program: AI & ML Microdegree', size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)

add_paragraph('Date: 23/04/2026', size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=20)

doc.add_page_break()

# ═══════════════════════════════════════════
# SUPERVISOR'S DECLARATION
# ═══════════════════════════════════════════

add_paragraph("SUPERVISOR'S DECLARATION", size=16, bold=True, space_after=24)

sup_text = (
    'I hereby confirm that the project proposal entitled "Earthquake Building '
    'Damage Prediction Using Machine Learning Algorithms" was carried out under '
    'my supervision and guidance. The work presented in this proposal meets the '
    'academic standards and requirements of the AI/ML Microdegree Program.'
)
add_paragraph(sup_text, size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=60)

add_paragraph('___________________________', size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=40, space_after=2)
add_paragraph('Mr. Diwash Sapkota', size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
add_paragraph('Designation: _______________', size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
add_paragraph('Code for Change, Pokhara', size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
add_paragraph('Date: _______________', size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)

doc.add_page_break()

# ═══════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════

add_paragraph('ABSTRACT', size=16, bold=True, space_after=24)

abs_p1 = (
    'Nepal is one of the most earthquake-prone countries in the world. '
    'The 2015 Gorkha earthquake alone destroyed over 600,000 buildings and '
    'caused massive loss of life. Even today, many buildings across the country '
    'remain vulnerable, and there is no easy way to figure out which ones '
    'are most at risk. This is a serious problem because being able to identify '
    'weak structures before an earthquake hits could save lives and help the '
    'government plan better.'
)
add_paragraph(abs_p1, size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)

abs_p2 = (
    'In this project, we try to solve this problem by building a machine learning '
    'system that can predict the level of damage a building might face during an '
    'earthquake. We use a publicly available dataset from DrivenData that contains '
    'information about more than 260,000 real buildings affected by the 2015 Nepal '
    'earthquake. The dataset includes details like the number of floors, building '
    'age, type of foundation, roof material, and location.'
)
add_paragraph(abs_p2, size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)

abs_p3 = (
    'We plan to train and compare several machine learning models including '
    'Decision Tree, K-Nearest Neighbors, Random Forest, and XGBoost to see which '
    'one gives the best results for predicting damage grades. The damage is '
    'classified into three levels: low, medium, and severe. We also plan to build '
    'a simple web application where users can enter building details and get a '
    'prediction about how vulnerable that building might be.'
)
add_paragraph(abs_p3, size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=20)

# Keywords
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.5
run_b = p.add_run('Keywords: ')
run_b.font.name = 'Times New Roman'
run_b.font.size = Pt(12)
run_b.font.bold = True
run_t = p.add_run(
    'Earthquake Damage Prediction, Machine Learning, Random Forest, '
    'XGBoost, Nepal, Classification, Disaster Preparedness'
)
run_t.font.name = 'Times New Roman'
run_t.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════

output_path = r'd:\Tele-Health\Earthquake_Damage_Prediction_Proposal.docx'
doc.save(output_path)
print(f'Proposal saved to: {output_path}')

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Inches(1.5)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

def remove_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        borders.append(border)
    tblPr.append(borders)

def add_para(text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold
    return p

def add_toc_entry(table, row, heading, page, indent=0, bold=False, size=12):
    cell0 = table.rows[row].cells[0]
    cell1 = table.rows[row].cells[1]
    
    cell0.text = ''
    p = cell0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    if indent > 0:
        p.paragraph_format.left_indent = Inches(indent * 0.4)
    r = p.add_run(heading)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold
    
    cell1.text = ''
    p2 = cell1.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(3)
    p2.paragraph_format.space_after = Pt(3)
    p2.paragraph_format.line_spacing = 1.5
    r2 = p2.add_run(str(page))
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(size)
    r2.font.bold = bold

# ═══════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════

add_para('TABLE OF CONTENTS', size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=24)

# Actual page numbers from the document!
# Front matter uses Roman numerals (offset: cover=page1, so Ack=page2=i, etc.)
# Main content: Introduction is page 6 in absolute, but needs to be "1"
# So Arabic offset = 6-1 = 5 (subtract 5 from absolute page)

# FRONT MATTER (Roman numerals)
# Cover = page 1 (no number)
# Acknowledgement = page 2 → i
# Student's Declaration = page 3 → ii  
# Abstract = page 4 → iii
# Table of Contents = will be inserted → iv
# List of Figures/Tables = → v
# Abbreviations = page 5 → vi

# MAIN CONTENT (Arabic - page 6 = 1)
# Introduction = page 6 → 1
# Problem Statement = page 8 → 3
# etc.

main_offset = 5  # subtract this from absolute page to get Arabic number

toc_entries = [
    # (heading, page_display, indent, bold)
    ('Acknowledgement', 'i', 0, False),
    ('Student\u2019s Declaration', 'ii', 0, False),
    ('Abstract', 'iii', 0, False),
    ('Table of Contents', 'iv', 0, False),
    ('List of Figures / List of Tables', 'v', 0, False),
    ('Abbreviations', 'vi', 0, False),
    ('', '', 0, False),  # spacer
    ('1. Introduction', str(6 - main_offset), 0, True),
    ('2. Problem Statement', str(8 - main_offset), 0, True),
    ('3. Objectives', str(9 - main_offset), 0, True),
    ('4. Literature Review', str(10 - main_offset), 0, True),
    ('5. Methodology', str(12 - main_offset), 0, True),
    ('5.1 Tools and Technologies', str(12 - main_offset), 1, False),
    ('5.2 Dataset', str(12 - main_offset), 1, False),
    ('5.3 Model Selection Rationale', str(13 - main_offset), 1, False),
    ('5.4 Training Approach', str(13 - main_offset), 1, False),
    ('5.5 Model Pipeline Diagram', str(15 - main_offset), 1, False),
    ('6. Requirement Document', str(16 - main_offset), 0, True),
    ('6.1 Functional Requirements', str(16 - main_offset), 1, False),
    ('6.2 Non-Functional Requirements', str(17 - main_offset), 1, False),
    ('6.3 Technology Stack', str(17 - main_offset), 1, False),
    ('7. System Analysis and Design', str(18 - main_offset), 0, True),
    ('7.1 System Architecture Diagram', str(18 - main_offset), 1, False),
    ('7.2 Data Flow Diagram', str(19 - main_offset), 1, False),
    ('7.3 Process Flowchart', str(20 - main_offset), 1, False),
    ('7.4 Model Pipeline Diagram', str(20 - main_offset), 1, False),
    ('7.5 Note on ER Diagram', str(21 - main_offset), 1, False),
    ('8. Development Plan', str(21 - main_offset), 0, True),
    ('9. Testing Strategy', str(23 - main_offset), 0, True),
    ('9.1 Test Types', str(23 - main_offset), 1, False),
    ('9.2 Sample Test Cases', str(23 - main_offset), 1, False),
    ('9.3 Validation Plan', str(24 - main_offset), 1, False),
    ('10. Expected Results', str(25 - main_offset), 0, True),
    ('11. Future Enhancements', str(26 - main_offset), 0, True),
    ('12. Conclusion', str(27 - main_offset), 0, True),
    ('13. Annex / Appendix', str(28 - main_offset), 0, True),
    ('14. Gantt Chart', str(30 - main_offset), 0, True),
    ('References', str(31 - main_offset), 0, True),
]

table = doc.add_table(rows=len(toc_entries), cols=2)
remove_borders(table)

for i, (heading, page, indent, bold) in enumerate(toc_entries):
    if heading == '':
        continue
    add_toc_entry(table, i, heading, page, indent, bold)

table.columns[0].width = Inches(5.0)
table.columns[1].width = Inches(0.7)

# Add dotted leader effect with tab stops would be ideal
# but python-docx tab leaders are limited, so we use the clean format

# ═══════════════════════════════════════
# LIST OF FIGURES
# ═══════════════════════════════════════

doc.add_page_break()

add_para('LIST OF FIGURES', size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=24)

figures = [
    ('Figure 1:', 'End-to-End Machine Learning Pipeline for Earthquake Building Damage Prediction', str(15 - main_offset)),
    ('Figure 2:', 'System Architecture Diagram', str(18 - main_offset)),
    ('Figure 3:', 'Data Flow Diagram \u2014 Level 0 (Context) and Level 1', str(19 - main_offset)),
    ('Figure 4:', 'Process Flowchart for Earthquake Damage Prediction System', str(20 - main_offset)),
    ('Figure 5:', 'Project Gantt Chart \u2014 Earthquake Building Damage Prediction', str(30 - main_offset)),
]

fig_table = doc.add_table(rows=len(figures), cols=3)
remove_borders(fig_table)

for i, (num, caption, page) in enumerate(figures):
    c0 = fig_table.rows[i].cells[0]
    c1 = fig_table.rows[i].cells[1]
    c2 = fig_table.rows[i].cells[2]
    
    for c, text, bold, align in [
        (c0, num, True, WD_ALIGN_PARAGRAPH.LEFT),
        (c1, caption, False, WD_ALIGN_PARAGRAPH.LEFT),
        (c2, page, False, WD_ALIGN_PARAGRAPH.RIGHT),
    ]:
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.bold = bold

fig_table.columns[0].width = Inches(0.9)
fig_table.columns[1].width = Inches(4.3)
fig_table.columns[2].width = Inches(0.5)

add_para('', after=24)

# ═══════════════════════════════════════
# LIST OF TABLES
# ═══════════════════════════════════════

add_para('LIST OF TABLES', size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=24)

tables_list = [
    ('Table 1:', 'Technology Stack', str(17 - main_offset)),
    ('Table 2:', 'Development Plan \u2014 Phases, Activities, and Timeline', str(21 - main_offset)),
    ('Table 3:', 'Sample Test Cases', str(23 - main_offset)),
    ('Table 4:', 'Expected Model Performance', str(25 - main_offset)),
    ('Table 5:', 'Dataset Information', str(28 - main_offset)),
    ('Table 6:', 'Key Python Libraries and Versions', str(28 - main_offset)),
    ('Table 7:', 'Sample Features from Dataset', str(29 - main_offset)),
]

tbl_table = doc.add_table(rows=len(tables_list), cols=3)
remove_borders(tbl_table)

for i, (num, caption, page) in enumerate(tables_list):
    c0 = tbl_table.rows[i].cells[0]
    c1 = tbl_table.rows[i].cells[1]
    c2 = tbl_table.rows[i].cells[2]
    
    for c, text, bold, align in [
        (c0, num, True, WD_ALIGN_PARAGRAPH.LEFT),
        (c1, caption, False, WD_ALIGN_PARAGRAPH.LEFT),
        (c2, page, False, WD_ALIGN_PARAGRAPH.RIGHT),
    ]:
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.bold = bold

tbl_table.columns[0].width = Inches(0.9)
tbl_table.columns[1].width = Inches(4.3)
tbl_table.columns[2].width = Inches(0.5)

path = r'd:\Tele-Health\TOC_Final.docx'
doc.save(path)
print(f'Saved: {path}')
print()
print('Page mapping used:')
print('  Cover = no number')
print('  Acknowledgement = i (absolute page 2)')
print('  Student Declaration = ii (absolute page 3)')
print('  Abstract = iii (absolute page 4)')
print('  TOC = iv')
print('  List of Figures/Tables = v')
print('  Abbreviations = vi (absolute page 5)')
print('  Introduction = 1 (absolute page 6)')
print(f'  References = {31 - main_offset} (absolute page 31)')

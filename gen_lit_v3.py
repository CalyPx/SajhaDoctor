from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Inches(1.5)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

def add_para(text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=8, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    return p

def add_mixed(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    for text, bold, italic in parts:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.bold = bold
        r.font.italic = italic
    return p

# ── HEADING ──
add_para('4. BACKGROUND STUDY / LITERATURE REVIEW', size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=12)

# ── Intro paragraph ──
add_para(
    'Earthquake damage prediction has become an active research area, '
    'especially after large-scale post-earthquake datasets became publicly '
    'available. Traditional vulnerability assessment relies on manual '
    'inspection, which is slow, expensive, and difficult to scale. This '
    'section reviews five studies that directly relate to the machine '
    'learning approach taken in this project.'
)

# ── Study 1: Ghimire et al. (2022) ──
add_mixed([
    ('Ghimire, Gu\u00e9guen, Giffard-Roisin, and Schorlemmer (2022)', False, False),
    (' analyzed 762,106 buildings from the Nepal Building Damage Portfolio '
     'across 11 districts, combined with macro-seismic intensity values from '
     'the USGS ShakeMap. They tested six models and found that Random Forest '
     'Regression performed best, achieving an accuracy of 0.68 for three-class '
     'damage classification and 0.49 for the full five-grade scale. Their '
     'feature importance analysis showed mud-mortar stone construction material '
     '(32%) and macro-seismic intensity (31%) as the most significant '
     'predictors. Even with only four basic building features, the model '
     'maintained a reasonable accuracy of 0.64.', False, False),
])

# ── Study 2: Bhatta & Dang (2023) ──
add_mixed([
    ('Bhatta and Dang (2023)', False, False),
    (' trained five models \u2014 KNN, Random Forest, Decision Tree, SVM, and '
     'ANN \u2014 on simulation data from 35 RC buildings and validated them on '
     '67 real RC buildings damaged during the 2015 Nepal earthquake. Random '
     'Forest achieved the highest testing accuracy of 74.62%, followed by ANN '
     '(70.14%), SVM (68.65%), Decision Tree (64.17%), and KNN (62.68%). '
     'Construction years, plinth area, and peak ground acceleration were '
     'identified as the most important features.', False, False),
])

# ── Study 3: Adi et al. (2020) ──
add_mixed([
    ('Adi, Adishesha, Bharadwaj, and Narayan (2020)', False, False),
    (' worked with the same DrivenData Nepal earthquake dataset used in this '
     'project. Their Random Forest Classifier achieved an F1-score of 72.95%, '
     'which improved to 74.42% using Gradient Boosting with hyperparameter '
     'tuning. This study serves as a direct benchmark since we use the same '
     'dataset and target variable.', False, False),
])

# ── Study 4: Poudyal & Shakya (2025) ──
add_mixed([
    ('Poudyal and Shakya (2025)', False, False),
    (' also used the DrivenData dataset and compared resampling techniques '
     'for class imbalance. Random Forest with SMOTE achieved a macro F1-score '
     'of 0.67 and accuracy of 82.1%. Both Random Forest and Gradient Boosting '
     'achieved the highest ROC-AUC of 0.89. They found Grade 2 (medium damage) '
     'was the hardest class to predict across all models due to overlapping '
     'features.', False, False),
])

# ── Study 5: Onur et al. (2022) ──
add_mixed([
    ('Onur, Atici, Ye\u015fil\u0131rmak, and Hanc\u0131lar (2022)', False, False),
    (' studied 8,071 buildings damaged during the 2020 Elaz\u0131\u011f '
     'earthquake (Mw 6.8) in Turkey. Random Forest achieved the best accuracy '
     'of 85% and F1-score of 68% after applying Tomek\u2019s link '
     'under-sampling, confirming that Random Forest generalizes well across '
     'different seismic contexts.', False, False),
])

# ── Summary paragraph ──
add_para(
    'These studies consistently show that ensemble methods like Random Forest '
    'and Gradient Boosting outperform simpler models, with F1-scores ranging '
    'from 0.67 to 0.74 on the Nepal dataset. Class imbalance remains a shared '
    'challenge. However, none of these studies produce a deployable tool for '
    'local authorities. This project addresses that gap by combining model '
    'evaluation with a FastAPI web application for practical earthquake '
    'preparedness planning.'
)

# ══════════════════════════
# REFERENCES
# ══════════════════════════
add_para('References', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
         before=16, after=10)

refs = [
    ('Adi, S. P., Adishesha, V. B., Bharadwaj, K. V., & Narayan, A. (2020). '
     'Earthquake damage prediction using Random Forest and Gradient Boosting '
     'Classifier. ',
     'American Journal of Biological and Environmental Statistics, 6',
     '(3), 58\u201363. https://doi.org/10.11648/j.ajbes.20200603.14'),

    ('Bhatta, S., & Dang, J. (2023). Seismic damage prediction of RC '
     'buildings using machine learning. ',
     'Earthquake Engineering & Structural Dynamics, 52',
     '(11), 3504\u20133527. https://doi.org/10.1002/eqe.3907'),

    ('Ghimire, S., Gu\u00e9guen, P., Giffard-Roisin, S., & Schorlemmer, D. '
     '(2022). Testing machine learning models for seismic damage prediction '
     'at a regional scale using building-damage dataset compiled after the '
     '2015 Gorkha Nepal earthquake. ',
     'Earthquake Spectra, 38',
     '(4), 2970\u20132993. https://doi.org/10.1177/87552930221106495'),

    ('Onur, U., Atici, A. T., Ye\u015fil\u0131rmak, F. C., & '
     'Hanc\u0131lar, U. (2022). Classifying and predicting earthquake damage '
     'by using machine learning after the 2020 Elaz\u0131\u011f, Turkey, '
     'earthquake. In ',
     'Proceedings of the 3rd European Conference on Earthquake Engineering '
     '& Seismology (3ECEES)',
     ', Bucharest, Romania.'),

    ('Poudyal, B., & Shakya, M. (2025). Enhancing earthquake preparedness '
     'in Nepal through machine learning-based damage prediction models. ',
     'Journal of Future Artificial Intelligence and Technologies, 2',
     '(3), 479\u2013492.'),
]

for normal_before, italic_part, normal_after in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)

    r1 = p.add_run(normal_before)
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(11)

    r2 = p.add_run(italic_part)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(11)
    r2.font.italic = True

    r3 = p.add_run(normal_after)
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(11)

path = r'd:\Tele-Health\Lit_Review_Condensed.docx'
doc.save(path)
print(f'Saved: {path}')

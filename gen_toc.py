from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Inches(1.5)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

def add_para(text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=8, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    return p

def set_cell(cell, text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold

def toc_row(table, row_idx, heading, page, indent=0, bold=False):
    cell0 = table.rows[row_idx].cells[0]
    cell1 = table.rows[row_idx].cells[1]
    cell0.text = ''
    p = cell0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    if indent > 0:
        p.paragraph_format.left_indent = Inches(indent * 0.4)
    r = p.add_run(heading)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.bold = bold
    set_cell(cell1, str(page), size=12, align=WD_ALIGN_PARAGRAPH.RIGHT)


# ═══════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════

add_para('TABLE OF CONTENTS', size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=20)

toc_entries = [
    ('Acknowledgement', 'i', 0, False),
    ("Student\u2019s Declaration", 'ii', 0, False),
    ('Abstract', 'iii', 0, False),
    ('Table of Contents', 'iv', 0, False),
    ('List of Figures', 'v', 0, False),
    ('List of Tables', 'v', 0, False),
    ('Abbreviations', 'vi', 0, False),
    ('', '', 0, False),
    ('1. Introduction', '1', 0, True),
    ('2. Problem Statement', '3', 0, True),
    ('3. Objectives', '4', 0, True),
    ('4. Background Study / Literature Review', '5', 0, True),
    ('5. Methodology', '7', 0, True),
    ('5.1 Tools and Technologies', '7', 1, False),
    ('5.2 Dataset', '7', 1, False),
    ('5.3 Model Selection Rationale', '8', 1, False),
    ('5.4 Training Approach', '8', 1, False),
    ('5.5 Model Pipeline Diagram', '9', 1, False),
    ('6. Requirement Document', '10', 0, True),
    ('6.1 Functional Requirements', '10', 1, False),
    ('6.2 Non-Functional Requirements', '11', 1, False),
    ('6.3 Technology Stack', '11', 1, False),
    ('7. System Analysis and Design', '12', 0, True),
    ('7.1 System Architecture Diagram', '12', 1, False),
    ('7.2 Data Flow Diagram', '13', 1, False),
    ('7.3 Process Flowchart', '14', 1, False),
    ('7.4 Model Pipeline Diagram', '14', 1, False),
    ('7.5 Note on ER Diagram', '14', 1, False),
    ('8. Development Plan', '15', 0, True),
    ('9. Testing Strategy', '16', 0, True),
    ('9.1 Test Types', '16', 1, False),
    ('9.2 Sample Test Cases', '16', 1, False),
    ('9.3 Validation Plan', '17', 1, False),
    ('10. Expected Results', '18', 0, True),
    ('11. Future Enhancements', '19', 0, True),
    ('12. Conclusion', '20', 0, True),
    ('13. Annex / Appendix', '21', 0, True),
    ('14. Gantt Chart', '23', 0, True),
    ('References', '24', 0, True),
]

# Create borderless table for TOC
toc_table = doc.add_table(rows=len(toc_entries), cols=2)

# Remove borders
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def remove_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        borders.append(border)
    tblPr.append(borders)

remove_borders(toc_table)

for i, (heading, page, indent, bold) in enumerate(toc_entries):
    if heading == '':
        continue
    toc_row(toc_table, i, heading, page, indent, bold)

toc_table.columns[0].width = Inches(5.0)
toc_table.columns[1].width = Inches(0.7)

# ═══════════════════════════════════════
# LIST OF FIGURES
# ═══════════════════════════════════════

doc.add_page_break()

add_para('LIST OF FIGURES', size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=20)

figures = [
    ('Figure 1', 'End-to-End Machine Learning Pipeline', '9'),
    ('Figure 2', 'System Architecture Diagram', '12'),
    ('Figure 3', 'Data Flow Diagram \u2014 Level 0 and Level 1', '13'),
    ('Figure 4', 'Process Flowchart for Earthquake Damage Prediction System', '14'),
    ('Figure 5', 'Project Gantt Chart', '23'),
]

fig_table = doc.add_table(rows=len(figures), cols=3)
remove_borders(fig_table)

for i, (num, caption, page) in enumerate(figures):
    set_cell(fig_table.rows[i].cells[0], num, bold=True, size=12)
    set_cell(fig_table.rows[i].cells[1], caption, size=12)
    set_cell(fig_table.rows[i].cells[2], page, size=12,
             align=WD_ALIGN_PARAGRAPH.RIGHT)

fig_table.columns[0].width = Inches(0.9)
fig_table.columns[1].width = Inches(4.3)
fig_table.columns[2].width = Inches(0.5)

add_para('', after=20)

# ═══════════════════════════════════════
# LIST OF TABLES
# ═══════════════════════════════════════

add_para('LIST OF TABLES', size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=20)

tables_list = [
    ('Table 1', 'Technology Stack', '11'),
    ('Table 2', 'Development Plan \u2014 Phases, Activities, and Timeline', '15'),
    ('Table 3', 'Sample Test Cases', '16'),
    ('Table 4', 'Expected Model Performance', '18'),
    ('Table 5', 'Dataset Information', '21'),
    ('Table 6', 'Key Python Libraries and Versions', '21'),
    ('Table 7', 'Sample Features from Dataset', '22'),
]

tbl_table = doc.add_table(rows=len(tables_list), cols=3)
remove_borders(tbl_table)

for i, (num, caption, page) in enumerate(tables_list):
    set_cell(tbl_table.rows[i].cells[0], num, bold=True, size=12)
    set_cell(tbl_table.rows[i].cells[1], caption, size=12)
    set_cell(tbl_table.rows[i].cells[2], page, size=12,
             align=WD_ALIGN_PARAGRAPH.RIGHT)

tbl_table.columns[0].width = Inches(0.9)
tbl_table.columns[1].width = Inches(4.3)
tbl_table.columns[2].width = Inches(0.5)

# ═══════════════════════════════════════
# ABBREVIATIONS
# ═══════════════════════════════════════

doc.add_page_break()

add_para('ABBREVIATIONS', size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=20)

abbreviations = [
    ('AI', 'Artificial Intelligence'),
    ('ANN', 'Artificial Neural Network'),
    ('APA', 'American Psychological Association'),
    ('API', 'Application Programming Interface'),
    ('CSV', 'Comma-Separated Values'),
    ('DFD', 'Data Flow Diagram'),
    ('DT', 'Decision Tree'),
    ('EDA', 'Exploratory Data Analysis'),
    ('ER', 'Entity-Relationship'),
    ('F1', 'F1-Score (Harmonic Mean of Precision and Recall)'),
    ('KNN', 'K-Nearest Neighbors'),
    ('ML', 'Machine Learning'),
    ('Mw', 'Moment Magnitude'),
    ('NBDP', 'Nepal Building Damage Portfolio'),
    ('RC', 'Reinforced Concrete'),
    ('RCC', 'Reinforced Cement Concrete'),
    ('REST', 'Representational State Transfer'),
    ('RF', 'Random Forest'),
    ('ROC-AUC', 'Receiver Operating Characteristic \u2013 Area Under Curve'),
    ('SMOTE', 'Synthetic Minority Over-sampling Technique'),
    ('SVM', 'Support Vector Machine'),
    ('USGS', 'United States Geological Survey'),
    ('VS30', 'Average Shear Wave Velocity in Top 30 Meters'),
    ('XGBoost', 'Extreme Gradient Boosting'),
]

abbr_table = doc.add_table(rows=len(abbreviations), cols=2)
remove_borders(abbr_table)

for i, (abbr, full) in enumerate(abbreviations):
    set_cell(abbr_table.rows[i].cells[0], abbr, bold=True, size=12)
    set_cell(abbr_table.rows[i].cells[1], full, size=12)

abbr_table.columns[0].width = Inches(1.3)
abbr_table.columns[1].width = Inches(4.5)

path = r'd:\Tele-Health\TOC_Figures_Abbreviations.docx'
doc.save(path)
print(f'Saved: {path}')

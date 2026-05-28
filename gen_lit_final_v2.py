from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Inches(1.5)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

def add_mixed(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=10):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    for text, bold, italic in parts:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.bold = bold
        r.font.italic = italic
    return p

def add_para(text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=10, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    return p

# ════════════════════════════════════
# HEADING
# ════════════════════════════════════

add_para('4. BACKGROUND STUDY / LITERATURE REVIEW', size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=18)

# ── Intro ──
add_para(
    'Earthquake damage prediction has become an active area of research '
    'in recent years, especially after large-scale post-earthquake datasets '
    'became publicly available. Traditional methods of assessing building '
    'vulnerability rely on manual inspection by engineers, which is slow, '
    'expensive, and hard to scale across entire regions. Several researchers '
    'have explored machine learning as a faster and more practical alternative. '
    'This section reviews key studies that are directly related to the approach '
    'taken in this project.',
    after=14
)

# ══════════════════════════════════════
# STUDY 1 — Ghimire et al. (2022)
# ══════════════════════════════════════

add_mixed([
    ('Ghimire, Guéguen, Giffard-Roisin, and Schorlemmer (2022) ', False, False),
    ('analyzed building damage data from the 2015 Gorkha Nepal earthquake at '
     'a regional scale. They used the Nepal Building Damage Portfolio (NBDP), '
     'which contained information on 762,106 buildings across 11 districts, and '
     'combined it with macro-seismic intensity values from the USGS ShakeMap tool. '
     'They tested six machine learning models including Linear Regression, Support '
     'Vector Regression, Gradient Boosting, and Random Forest. Among all the models '
     'tested, the Random Forest Regression model provided the best damage predictions. '
     'When buildings were classified into three damage categories (slight, medium, '
     'and heavy), the model achieved an accuracy of 0.68. With the full five-grade '
     'damage scale, the accuracy was 0.49. Their feature importance analysis showed '
     'that mud-mortar stone construction material had the highest importance score '
     '(32%), followed by macro-seismic intensity (31%). They also found that using '
     'only four basic building features (number of stories, age, height, and plinth '
     'area) along with ground-motion intensity still produced a reasonable accuracy '
     'of 0.64 for three-class classification. This study is relevant to our project '
     'because we also rely on building attributes from the same earthquake event '
     'and use Random Forest as one of our candidate models.', False, False),
], after=14)

# ══════════════════════════════════════
# STUDY 2 — Bhatta & Dang (2023)
# ══════════════════════════════════════

add_mixed([
    ('Bhatta and Dang (2023) ', False, False),
    ('focused specifically on reinforced concrete (RC) buildings and presented '
     'a machine learning-based approach for seismic damage prediction. They '
     'trained their models using simulation data generated from nonlinear time '
     'history analysis (NLTHA) of 35 RC buildings subjected to 11 earthquake '
     'waves. Five machine learning models were used: K-Nearest Neighbor, Random '
     'Forest, Decision Tree, Support Vector Machine, and Artificial Neural Network. '
     'On the training dataset, Random Forest achieved the highest accuracy of '
     '99.83%. The key contribution of their study was that they validated these '
     'models using real earthquake damage data from 67 RC buildings collected '
     'after the 2015 Nepal earthquake. On this real-world testing dataset, Random '
     'Forest again achieved the highest overall accuracy of 74.62%, followed by '
     'ANN at 70.14%, SVM at 68.65%, Decision Tree at 64.17%, and KNN at 62.68%. '
     'Their feature importance analysis identified construction years, plinth area, '
     'and peak ground acceleration as the most significant contributors to damage '
     'prediction. This study is directly relevant to our project because the same '
     'four model types — Decision Tree, KNN, Random Forest, and XGBoost — are '
     'evaluated in our work.', False, False),
], after=14)

# ══════════════════════════════════════
# STUDY 3 — Adi et al. (2020)
# ══════════════════════════════════════

add_mixed([
    ('Adi, Adishesha, Bharadwaj, and Narayan (2020) ', False, False),
    ('worked directly with the DrivenData Nepal earthquake dataset, which is '
     'the same dataset used in this project. The data was collected through '
     'surveys conducted by the Kathmandu Living Labs and the Central Bureau of '
     'Statistics under the National Planning Commission Secretariat of Nepal. '
     'They used Random Forest Classifier and Gradient Boosting Classifier to '
     'predict damage grade severity across three levels. With parameter tuning, '
     'the Random Forest Classifier achieved a micro-averaged F1-score of 72.95%. '
     'They then applied winsorization to handle outliers, which improved the '
     'F1-score to 74.33% when combined with Gradient Boosting. Finally, using '
     'only hyperparameter tuning with Gradient Boosting (without winsorization), '
     'they achieved their best F1-score of 74.42%. This study serves as a direct '
     'benchmark for our project since we are working with the same dataset and '
     'target variable.', False, False),
], after=14)

# ══════════════════════════════════════
# STUDY 4 — Poudyal & Shakya (2025)
# ══════════════════════════════════════

add_mixed([
    ('Poudyal and Shakya (2025) ', False, False),
    ('also used the DrivenData dataset from the 2015 Gorkha earthquake and '
     'tested six machine learning models: Random Forest, Multilayer Perceptron, '
     'Support Vector Machine, Gradient Boosting, and XGBoost. A key contribution '
     'of their study was a comparative analysis of resampling techniques for '
     'handling class imbalance. They tested three scenarios: no oversampling, '
     'Random Oversampling, and SMOTE. The Random Forest model with SMOTE '
     'oversampling achieved the best macro F1-score of 0.67 with an accuracy of '
     '82.1%, while Random Oversampling achieved a macro F1-score of 0.65 with an '
     'accuracy of 81.5%. Both Random Forest and Gradient Boosting achieved the '
     'highest macro-average ROC-AUC of 0.89. They also found that Grade 2 '
     '(medium damage) was consistently the most difficult class to predict across '
     'all models due to overlapping features with Grades 1 and 3. This finding '
     'is particularly relevant to our project because we also expect to encounter '
     'class imbalance, and their results inform our decision to use stratified '
     'sampling during train-test splitting.', False, False),
], after=14)

# ══════════════════════════════════════
# STUDY 5 — Onur et al. (2022)
# ══════════════════════════════════════

add_mixed([
    ('Onur, Atici, Yeşilırmak, and Hancılar (2022) ', False, False),
    ('studied building damage classification after the 2020 Elazığ earthquake '
     '(Mw 6.8) in Turkey. They used a dataset of 8,071 damaged buildings from '
     'the Merkez county of Elazığ, which were classified into moderate damage, '
     'extensive damage, and complete damage categories. Six machine learning '
     'algorithms were compared: Logistic Regression, Decision Tree, Random '
     'Forest, K-Nearest Neighbor, Support Vector Machine, and Naïve Bayes. The '
     'Random Forest algorithm achieved the best performance with a mean accuracy '
     'of 85% and an F1-score of 68% after applying Tomek\'s link under-sampling '
     'to handle class imbalance. This study from a different seismic context '
     'confirms that Random Forest consistently performs well for multi-class '
     'building damage prediction, supporting its inclusion as a candidate model '
     'in our project.', False, False),
], after=14)

# ── Summary ──
add_para(
    'The reviewed studies show a consistent pattern: ensemble methods like '
    'Random Forest and Gradient Boosting outperform simpler models for earthquake '
    'damage classification tasks. On the DrivenData Nepal dataset specifically, '
    'reported F1-scores fall in the range of 0.67 to 0.74 depending on '
    'preprocessing and tuning. Class imbalance, particularly the overrepresentation '
    'of medium-damage buildings, is a shared challenge across all studies. However, '
    'most existing studies treat the prediction task as a standalone research '
    'exercise and do not produce a deployable tool that can be used by local '
    'authorities or planners. This project aims to fill that gap by combining '
    'comparative model evaluation with the development of a web application '
    'for practical earthquake preparedness planning.',
    after=24
)

# ══════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════

add_para('References', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
         before=20, after=14)

refs = [
    'Adi, S. P., Adishesha, V. B., Bharadwaj, K. V., & Narayan, A. (2020). '
    'Earthquake damage prediction using Random Forest and Gradient Boosting '
    'Classifier. American Journal of Biological and Environmental Statistics, '
    '6(3), 58–63. https://doi.org/10.11648/j.ajbes.20200603.14',

    'Bhatta, S., & Dang, J. (2023). Seismic damage prediction of RC buildings '
    'using machine learning. Earthquake Engineering & Structural Dynamics, '
    '52(11), 3504–3527. https://doi.org/10.1002/eqe.3907',

    'Ghimire, S., Guéguen, P., Giffard-Roisin, S., & Schorlemmer, D. (2022). '
    'Testing machine learning models for seismic damage prediction at a '
    'regional scale using building-damage dataset compiled after the 2015 '
    'Gorkha Nepal earthquake. Earthquake Spectra, 38(4), 2970–2993. '
    'https://doi.org/10.1177/87552930221106495',

    'Onur, U., Atici, A. T., Yeşilırmak, F. C., & Hancılar, U. (2022). '
    'Classifying and predicting earthquake damage by using machine learning '
    'after the 2020 Elazığ, Turkey, earthquake. In Proceedings of the 3rd '
    'European Conference on Earthquake Engineering & Seismology (3ECEES), '
    'Bucharest, Romania.',

    'Poudyal, B., & Shakya, M. (2025). Enhancing earthquake preparedness in '
    'Nepal through machine learning-based damage prediction models. Journal '
    'of Future Artificial Intelligence and Technologies, 2(3), 479–492.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    r = p.add_run(ref)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)

# SAVE
path = r'd:\Tele-Health\Lit_Review_v3.docx'
doc.save(path)
print(f'Saved: {path}')

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Inches(1.5)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

def add_para(text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold
    return p

def set_cell(cell, text, bold=False, size=11):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold

# Heading
add_para('10. EXPECTED RESULTS', size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=12)

add_para(
    'Based on the literature reviewed and the nature of the dataset, '
    'we expect the following outcomes from this project:'
)

add_para(
    'Random Forest is expected to deliver the best overall performance '
    'among the four models. Previous studies on the same DrivenData dataset '
    'have reported F1-scores between 0.67 and 0.74 for similar classification '
    'tasks (Adi et al., 2020; Poudyal & Shakya, 2025). We anticipate our '
    'Random Forest model to achieve a macro-averaged F1-score in the range of '
    '0.70 to 0.74 after hyperparameter tuning. XGBoost is expected to perform '
    'close to Random Forest, while Decision Tree and KNN are likely to fall '
    'behind due to their simpler learning mechanisms.'
)

add_para(
    'Grade 2 (medium damage) is expected to be the hardest class to predict '
    'accurately. This has been a consistent finding across all reviewed studies '
    'because Grade 2 buildings share overlapping features with both Grade 1 '
    'and Grade 3. We expect the models to perform best on Grade 1 and Grade 3, '
    'with Grade 2 pulling down the overall macro F1-score.'
)

add_para(
    'The table below summarizes the anticipated performance range for each model:'
)

# Table
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'

headers = ['Model', 'Expected F1-Score', 'Expected Accuracy']
for i, h in enumerate(headers):
    set_cell(table.rows[0].cells[i], h, bold=True)

data = [
    ('Decision Tree', '0.60 \u2013 0.65', '65% \u2013 70%'),
    ('K-Nearest Neighbors', '0.58 \u2013 0.63', '62% \u2013 68%'),
    ('Random Forest', '0.70 \u2013 0.74', '72% \u2013 78%'),
    ('XGBoost', '0.68 \u2013 0.73', '70% \u2013 76%'),
]

for row_idx, (model, f1, acc) in enumerate(data, 1):
    set_cell(table.rows[row_idx].cells[0], model)
    set_cell(table.rows[row_idx].cells[1], f1)
    set_cell(table.rows[row_idx].cells[2], acc)

for row in table.rows:
    row.cells[0].width = Inches(2.0)
    row.cells[1].width = Inches(1.8)
    row.cells[2].width = Inches(1.8)

add_para('', after=4)

add_para(
    'On the deployment side, we expect the FastAPI web application to return '
    'predictions within 2 seconds for any valid input. The web form should '
    'correctly handle all building feature inputs and display the predicted '
    'damage grade in a clear and understandable format. The overall goal is '
    'to produce a working prototype that demonstrates how machine learning '
    'can be used for practical earthquake preparedness in Nepal.'
)

path = r'd:\Tele-Health\Expected_Results.docx'
doc.save(path)
print(f'Saved: {path}')

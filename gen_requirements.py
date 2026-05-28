from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Inches(1.5)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

def add_para(text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=8, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    return p

def add_bullet(text, bold_prefix='', indent=0.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    if bold_prefix:
        rb = p.add_run('\u2022 ' + bold_prefix)
        rb.font.name = 'Times New Roman'
        rb.font.size = Pt(12)
        rb.font.bold = True
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
    else:
        r = p.add_run('\u2022 ' + text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
    return p

def add_image(img_path, caption, fig_num, width=5.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(img_path, width=Inches(width))
    add_para(f'Figure {fig_num}: {caption}', size=10, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=12)


# ══════════════════════════════════════════════
# SECTION 6: REQUIREMENT DOCUMENT
# ══════════════════════════════════════════════

add_para('6. REQUIREMENT DOCUMENT', size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=12)

# --- 6.1 Functional Requirements ---
add_para('6.1 Functional Requirements', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=6)

add_para(
    'The following functional requirements describe what the system '
    'should be able to do:'
)

func_reqs = [
    ('FR-01: ', 'The system shall load and merge the DrivenData earthquake '
     'dataset consisting of building features and damage grade labels.'),
    ('FR-02: ', 'The system shall perform data preprocessing including '
     'encoding categorical variables, handling missing values, and splitting '
     'data into training and testing sets.'),
    ('FR-03: ', 'The system shall train and evaluate four machine learning '
     'models: Decision Tree, K-Nearest Neighbor, Random Forest, and XGBoost.'),
    ('FR-04: ', 'The system shall display model comparison results including '
     'F1-score, accuracy, precision, recall, and confusion matrices.'),
    ('FR-05: ', 'The system shall provide a web-based form where users can '
     'input building characteristics such as number of floors, age, '
     'construction material, foundation type, and roof type.'),
    ('FR-06: ', 'The system shall predict and display the damage grade '
     '(Grade 1, 2, or 3) for the submitted building based on the '
     'best-performing trained model.'),
    ('FR-07: ', 'The system shall visualize feature importance rankings '
     'to show which building attributes contribute most to the prediction.'),
]

for prefix, desc in func_reqs:
    add_bullet(desc, bold_prefix=prefix)

# --- 6.2 Non-Functional Requirements ---
add_para('6.2 Non-Functional Requirements', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=6)

add_para(
    'The non-functional requirements define the quality attributes and '
    'constraints of the system:'
)

nonfunc_reqs = [
    ('NFR-01 (Performance): ', 'The prediction response time shall be '
     'under 2 seconds after the user submits the form.'),
    ('NFR-02 (Accuracy): ', 'The selected model shall achieve a '
     'micro-averaged F1-score of at least 0.70 on the test dataset.'),
    ('NFR-03 (Usability): ', 'The web interface shall be simple enough '
     'for a non-technical user to fill in building details and understand '
     'the prediction result.'),
    ('NFR-04 (Compatibility): ', 'The web application shall run on '
     'modern browsers including Chrome, Firefox, and Edge.'),
    ('NFR-05 (Portability): ', 'The system shall be deployable on a '
     'local machine without requiring cloud infrastructure.'),
    ('NFR-06 (Maintainability): ', 'The codebase shall be modular and '
     'documented so that individual components can be updated independently.'),
]

for prefix, desc in nonfunc_reqs:
    add_bullet(desc, bold_prefix=prefix)

# --- 6.3 Technology Stack ---
add_para('6.3 Technology Stack', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=6)

# Simple table for tech stack
table = doc.add_table(rows=10, cols=2)
table.style = 'Table Grid'

headers = ['Component', 'Technology']
data = [
    ('Programming Language', 'Python 3.x'),
    ('ML Libraries', 'scikit-learn, XGBoost'),
    ('Data Processing', 'pandas, NumPy'),
    ('Visualization', 'matplotlib, seaborn'),
    ('Backend Framework', 'FastAPI, uvicorn'),
    ('Frontend', 'HTML, CSS, JavaScript'),
    ('Model Serialization', 'joblib'),
    ('Development Environment', 'Jupyter Notebook'),
    ('Version Control', 'Git'),
]

for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    r = cell.paragraphs[0].add_run(h)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.font.bold = True

for row_idx, (comp, tech) in enumerate(data, 1):
    for col_idx, val in enumerate([comp, tech]):
        cell = table.rows[row_idx].cells[col_idx]
        cell.text = ''
        r = cell.paragraphs[0].add_run(val)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)

add_para('', after=4)  # spacer

# ══════════════════════════════════════════════
# SECTION 7: SYSTEM ANALYSIS & DESIGN
# ══════════════════════════════════════════════

doc.add_page_break()

add_para('7. SYSTEM ANALYSIS AND DESIGN', size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=12)

# --- 7.1 System Architecture ---
add_para('7.1 System Architecture Diagram', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=6)

add_para(
    'The system follows a three-layer architecture. The frontend layer '
    'handles user interaction through a web form where building details '
    'are entered. The backend layer, built with FastAPI, receives these '
    'inputs, validates them, and passes them to the prediction engine. '
    'The ML engine layer contains the trained model file which was saved '
    'during the training phase using joblib. The training pipeline itself '
    'runs separately in Jupyter Notebook and only needs to be executed '
    'once to produce the model file. Figure 1 shows the overall system '
    'architecture.'
)

arch_img = r'C:\Users\Gaming F16\.gemini\antigravity\brain\4ee70f2d-ccac-4272-98d6-3026c937c32e\system_architecture_1776922711614.png'
add_image(arch_img, 'System Architecture Diagram', 1)

# --- 7.2 Data Flow Diagram ---
add_para('7.2 Data Flow Diagram', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=6)

add_para(
    'The Data Flow Diagram shows how data moves through the system. '
    'At Level 0 (context level), the system interacts with two external '
    'entities: the User who provides building features and receives '
    'predictions, and the DrivenData platform which supplies the training '
    'dataset. At Level 1, the system is broken down into three main '
    'processes: data preprocessing, model training and evaluation, and the '
    'prediction engine. The trained model is stored as a data store that '
    'connects the training phase to the prediction phase. Figure 2 shows '
    'both DFD Level 0 and Level 1.'
)

dfd_img = r'C:\Users\Gaming F16\.gemini\antigravity\brain\4ee70f2d-ccac-4272-98d6-3026c937c32e\dfd_diagrams_1776922731436.png'
add_image(dfd_img, 'Data Flow Diagram \u2014 Level 0 (Context) and Level 1', 2)

# --- 7.3 Flowchart ---
add_para('7.3 Process Flowchart', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=6)

add_para(
    'The flowchart in Figure 3 shows the step-by-step logic of the '
    'system, starting from loading the dataset to returning the final '
    'prediction. After training and evaluating all four models, the system '
    'checks whether a satisfactory model has been found. If not, '
    'hyperparameters are tuned and the models are retrained. Once the best '
    'model is identified, it is saved and deployed through the FastAPI '
    'web application.'
)

flow_img = r'C:\Users\Gaming F16\.gemini\antigravity\brain\4ee70f2d-ccac-4272-98d6-3026c937c32e\flowchart_diagram_1776922951796.png'
add_image(flow_img, 'Process Flowchart for Earthquake Damage Prediction System', 3)

# --- 7.4 Model Pipeline ---
add_para('7.4 Model Pipeline Diagram', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=6)

add_para(
    'Figure 4 illustrates the end-to-end machine learning pipeline, '
    'from data ingestion through the DrivenData dataset to the final '
    'prediction output served via the web application. The pipeline '
    'covers data collection, exploratory analysis, preprocessing, feature '
    'selection, parallel training of four candidate models, evaluation, '
    'and deployment.'
)

pipeline_img = r'C:\Users\Gaming F16\.gemini\antigravity\brain\4ee70f2d-ccac-4272-98d6-3026c937c32e\ml_pipeline_diagram_1776918412604.png'
add_image(pipeline_img, 'End-to-End Machine Learning Pipeline', 4)

# --- Note about ER diagram ---
add_para('7.5 Note on ER Diagram', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=6)

add_para(
    'This project does not use a traditional relational database. The '
    'dataset is loaded directly from CSV files during training, and the '
    'trained model is stored as a serialized file. The web application '
    'processes each prediction request independently without storing '
    'user data in a database. Therefore, an Entity-Relationship diagram '
    'is not applicable to this project.'
)

# SAVE
path = r'd:\Tele-Health\Requirements_and_Diagrams.docx'
doc.save(path)
print(f'Saved: {path}')

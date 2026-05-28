from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page Setup — A4, margins per guidelines
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Inches(1.5)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

def add_para(text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
             before=0, after=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    return p

def add_justified(text, size=12, bold=False, before=0, after=12):
    return add_para(text, size, bold, WD_ALIGN_PARAGRAPH.JUSTIFY, before, after)

# ══════════════════════════════════════
# PAGE 1 — COVER PAGE
# ══════════════════════════════════════

add_para('', size=12, after=60)
add_para('Earthquake Building Damage Prediction', size=16, bold=True, after=0)
add_para('Using Machine Learning Algorithms', size=16, bold=True, after=24)
add_para('A Project Proposal', size=14, bold=True, before=30, after=12)
add_para('Submitted in partial fulfillment of the requirements', size=12, before=20, after=0)
add_para('for the completion of the', size=12, after=0)
add_para('Artificial Intelligence & Machine Learning Microdegree', size=12, bold=True, after=40)
add_para('Submitted by:', size=12, bold=True, before=40, after=8)
add_para('Rohit Poudel', size=12, after=2)
add_para('Enrollment No.: _______________', size=11, after=12)
add_para('Keshav KC', size=12, after=2)
add_para('Enrollment No.: _______________', size=11, after=30)
add_para('Supervised by:', size=12, bold=True, before=20, after=8)
add_para('Mr. Diwash Sapkota', size=12, after=30)
add_para('Code for Change, Pokhara', size=13, bold=True, before=30, after=40)
add_para('Submission Date: 23/04/2026', size=12, before=20)

doc.add_page_break()

# ══════════════════════════════════════
# PAGE 2 — ACKNOWLEDGEMENT
# ══════════════════════════════════════

add_para('ACKNOWLEDGEMENT', size=16, bold=True, after=24)

add_justified(
    'We would like to express our sincere gratitude to our supervisor, '
    'Mr. Diwash Sapkota, for his guidance and support throughout the '
    'development of this project proposal. His feedback and suggestions '
    'helped us shape our ideas into a proper plan.'
)

add_justified(
    'We are grateful to Code for Change, Pokhara for organizing the AI/ML '
    'Microdegree Program and giving us the platform to learn and work on '
    'real projects. We would also like to thank Gandaki University and '
    'National Innovation Centre ICT Lab for their academic support and for '
    'making this program possible. Our thanks also go to AI Community Nepal '
    '(AICN) for their collaboration in this initiative.'
)

add_justified(
    'We appreciate the open-source community and platforms like DrivenData '
    'for making the Nepal earthquake dataset publicly available, which forms '
    'the foundation of this project. Finally, we would like to thank our '
    'families and friends for their constant encouragement.'
)

add_para('Rohit Poudel', size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, before=40, after=2)
add_para('Keshav KC', size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, after=2)
add_para('April 2026', size=12, align=WD_ALIGN_PARAGRAPH.RIGHT)

doc.add_page_break()

# ══════════════════════════════════════
# PAGE 3 — STUDENT'S DECLARATION
# ══════════════════════════════════════

add_para("STUDENT'S DECLARATION", size=16, bold=True, after=24)

add_justified(
    'We hereby declare that we are the sole authors of this project proposal '
    'and that no sources other than those referenced herein have been used. '
    'The work presented is original and any resemblance to existing projects '
    'is purely coincidental. We acknowledge that academic dishonesty will '
    'result in disqualification.'
)

for name in ['Rohit Poudel', 'Keshav KC']:
    add_para('___________________________', size=12, align=WD_ALIGN_PARAGRAPH.LEFT, before=30, after=2)
    add_para(name, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, after=2)
    add_para('Enrollment No.: _______________', size=11, align=WD_ALIGN_PARAGRAPH.LEFT, after=2)
    add_para('Program: AI & ML Microdegree', size=11, align=WD_ALIGN_PARAGRAPH.LEFT, after=2)

add_para('Date: 23/04/2026', size=12, align=WD_ALIGN_PARAGRAPH.LEFT, before=20)

doc.add_page_break()

# ══════════════════════════════════════
# PAGE 4 — SUPERVISOR'S DECLARATION
# ══════════════════════════════════════

add_para("SUPERVISOR'S DECLARATION", size=16, bold=True, after=24)

add_justified(
    'I hereby confirm that the project proposal entitled "Earthquake Building '
    'Damage Prediction Using Machine Learning Algorithms" was carried out under '
    'my supervision and guidance. The work presented in this proposal meets the '
    'academic standards and requirements of the AI/ML Microdegree Program.'
)

add_para('___________________________', size=12, align=WD_ALIGN_PARAGRAPH.LEFT, before=50, after=2)
add_para('Mr. Diwash Sapkota', size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, after=2)
add_para('Designation: _______________', size=11, align=WD_ALIGN_PARAGRAPH.LEFT, after=2)
add_para('Code for Change, Pokhara', size=11, align=WD_ALIGN_PARAGRAPH.LEFT, after=2)
add_para('Date: _______________', size=11, align=WD_ALIGN_PARAGRAPH.LEFT)

doc.add_page_break()

# ══════════════════════════════════════
# PAGE 5 — ABSTRACT
# ══════════════════════════════════════

add_para('ABSTRACT', size=16, bold=True, after=24)

add_justified(
    'Nepal sits in one of the most seismically active regions in the world. '
    'The 2015 Gorkha earthquake caused widespread destruction, damaging over '
    '600,000 buildings and taking thousands of lives. A large number of '
    'buildings across the country still remain structurally vulnerable, but '
    'there is no practical system in place to assess which structures are '
    'most likely to suffer heavy damage in future earthquakes. Identifying '
    'high-risk buildings ahead of time could help the government and local '
    'authorities take preventive action and allocate resources more effectively.'
)

add_justified(
    'This project proposes a machine learning-based system for predicting '
    'the damage grade of buildings in the event of an earthquake. The system '
    'uses a publicly available dataset from DrivenData containing records of '
    'over 260,000 buildings affected by the 2015 Nepal earthquake. Building '
    'attributes such as age, number of floors, foundation type, roof material, '
    'and geographic location are used as input features. The target variable '
    'is the damage grade, classified into three levels: low, medium, and severe.'
)

add_justified(
    'Several classification algorithms including Decision Tree, K-Nearest '
    'Neighbors, Random Forest, and XGBoost are trained and compared based on '
    'accuracy and F1-score. The best-performing model is then deployed through '
    'a simple web interface built using Python and FastAPI. The expected outcome '
    'is a working prediction tool that achieves at least 70% classification '
    'accuracy and can be used for earthquake preparedness planning in Nepal.'
)

# Keywords with bold label
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.line_spacing = 1.5
run_b = p.add_run('Keywords: ')
run_b.font.name = 'Times New Roman'
run_b.font.size = Pt(12)
run_b.font.bold = True
run_t = p.add_run(
    'Earthquake Damage Prediction, Machine Learning, Classification, '
    'Random Forest, XGBoost, Nepal'
)
run_t.font.name = 'Times New Roman'
run_t.font.size = Pt(12)

# ══════════════════════════════════════
# SAVE
# ══════════════════════════════════════

path = r'd:\Tele-Health\Proposal_v2.docx'
doc.save(path)
print(f'Saved: {path}')

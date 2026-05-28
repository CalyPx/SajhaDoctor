from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Inches(1.5)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

def add_para(text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold
    return p

# ═══════════════════════════════
# SECTION 11: FUTURE ENHANCEMENTS
# ═══════════════════════════════

add_para('11. FUTURE ENHANCEMENTS', size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=12)

add_para(
    'While this project covers the core objective of predicting earthquake '
    'damage using machine learning, there are a number of things that could '
    'be added or improved if more time and resources were available.'
)

add_para(
    'One obvious next step would be to include geographic and soil data '
    'as additional features. Right now, the model only looks at building '
    'characteristics like age, floor count, and construction material. But '
    'the type of soil a building sits on and how close it is to a fault '
    'line also play a big role in how much damage it takes. Adding features '
    'like VS30 (soil shear wave velocity) or distance from the epicenter '
    'could help the model make better predictions, especially for buildings '
    'in areas where soil conditions vary a lot.'
)

add_para(
    'Another improvement would be to try deep learning models like '
    'Multilayer Perceptron or even more advanced architectures. The current '
    'project sticks to traditional machine learning algorithms because '
    'they are easier to interpret and faster to train, but deep learning '
    'might pick up on patterns that simpler models miss, particularly for '
    'Grade 2 buildings which are the hardest to classify correctly.'
)

add_para(
    'The web application itself could also be expanded. Right now it takes '
    'in details for one building at a time. A useful addition would be to '
    'allow batch predictions, where someone could upload a CSV file of '
    'hundreds of buildings and get damage predictions for all of them at '
    'once. This would be much more practical for local government offices '
    'or disaster management agencies that need to assess entire districts.'
)

add_para(
    'Finally, the system could be extended to include a mapping component. '
    'If geographic coordinates are available for each building, the '
    'predictions could be plotted on an interactive map showing which areas '
    'have the highest concentration of vulnerable buildings. This kind of '
    'visual output would be far more useful for planning and resource '
    'allocation than a simple table of numbers.'
)

# ═══════════════════════════════
# SECTION 12: CONCLUSION
# ═══════════════════════════════

doc.add_page_break()

add_para('12. CONCLUSION', size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=12)

add_para(
    'Nepal sits in one of the most earthquake-prone regions in the world, '
    'and the 2015 Gorkha earthquake made it painfully clear how vulnerable '
    'a large portion of the country\u2019s buildings really are. Manually '
    'inspecting millions of structures across 77 districts is simply not '
    'feasible with the resources currently available. This project proposes '
    'a machine learning-based approach as a practical alternative.'
)

add_para(
    'Using the publicly available DrivenData dataset of over 260,000 '
    'buildings, we plan to train and compare four classification models '
    '\u2014 Decision Tree, K-Nearest Neighbors, Random Forest, and '
    'XGBoost \u2014 to predict building damage grades. Based on what '
    'previous studies have shown, we expect Random Forest to perform best '
    'with an F1-score around 0.70 to 0.74. The best model will then be '
    'deployed through a simple web application built with FastAPI, allowing '
    'anyone to input building details and get a damage prediction without '
    'needing any technical background.'
)

add_para(
    'This is not meant to replace proper engineering assessments. What it '
    'can do is provide a quick, data-driven first screening that helps '
    'prioritize which buildings need urgent attention. If a local authority '
    'has limited engineers and a large number of buildings to check, a tool '
    'like this could help them decide where to start. That is the practical '
    'value this project aims to deliver.'
)

path = r'd:\Tele-Health\Future_Conclusion.docx'
doc.save(path)
print(f'Saved: {path}')

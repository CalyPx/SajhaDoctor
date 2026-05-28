from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ═══════════════════════════════════════
# PAGE 1: DEVELOPMENT PLAN (Portrait)
# ═══════════════════════════════════════

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Inches(1.5)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

def add_para(text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold
    return p

def set_cell(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold

def shade_cell(cell, color='4472C4'):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

# ── Development Plan Heading ──
add_para('8. DEVELOPMENT PLAN', size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=12)

add_para(
    'The project is planned over a 4-week period divided into five phases. '
    'The responsibilities are split between the two team members, though '
    'both collaborate on research and documentation. The table below '
    'outlines each phase with its activities, milestones, and timeline.'
)

# Dev Plan Table
dp_table = doc.add_table(rows=6, cols=5)
dp_table.style = 'Table Grid'

dp_headers = ['Phase', 'Activities', 'Milestone', 'Timeline', 'Responsibility']
for i, h in enumerate(dp_headers):
    set_cell(dp_table.rows[0].cells[i], h, bold=True, size=10)
    shade_cell(dp_table.rows[0].cells[i], 'D9E2F3')

dp_data = [
    ('Phase 1:\nResearch &\nPlanning',
     'Literature review, proposal writing, dataset study, tool setup',
     'Proposal approved, environment ready',
     'Week 1',
     'Both'),
    ('Phase 2:\nData Analysis &\nPreprocessing',
     'Download dataset, perform EDA, handle missing values, encode features, split data',
     'Clean dataset ready for training',
     'Week 1\u20132',
     'Rohit Poudel'),
    ('Phase 3:\nModel\nDevelopment',
     'Train DT, KNN, RF, and XGBoost. Hyperparameter tuning with GridSearchCV. Feature importance analysis',
     'All four models trained and evaluated',
     'Week 2\u20133',
     'Keshav KC'),
    ('Phase 4:\nDeployment',
     'Save best model with joblib, build FastAPI backend, develop web form frontend',
     'Working web application',
     'Week 3\u20134',
     'Rohit Poudel'),
    ('Phase 5:\nTesting &\nDocumentation',
     'Test web app, verify accuracy, write final report, prepare presentation',
     'Final report submitted, project demonstrated',
     'Week 4',
     'Both'),
]

for row_idx, (phase, act, mile, time, resp) in enumerate(dp_data, 1):
    set_cell(dp_table.rows[row_idx].cells[0], phase, bold=True, size=9)
    set_cell(dp_table.rows[row_idx].cells[1], act, size=9)
    set_cell(dp_table.rows[row_idx].cells[2], mile, size=9)
    set_cell(dp_table.rows[row_idx].cells[3], time, size=9)
    set_cell(dp_table.rows[row_idx].cells[4], resp, size=9)

# Column widths
for row in dp_table.rows:
    row.cells[0].width = Inches(1.0)
    row.cells[1].width = Inches(2.0)
    row.cells[2].width = Inches(1.5)
    row.cells[3].width = Inches(0.7)
    row.cells[4].width = Inches(1.0)

add_para('', after=6)

add_para('Responsibilities:', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=6)

add_para(
    'This project is carried out by two team members under the supervision '
    'of Mr. Diwash Sapkota. Rohit Poudel handles data preprocessing, EDA, '
    'and web application development. Keshav KC handles model training, '
    'hyperparameter tuning, and evaluation. Both members contribute to the '
    'research, documentation, and testing phases.'
)

# ═══════════════════════════════════════
# PAGE 2: GANTT CHART (Landscape)
# ═══════════════════════════════════════

new_section = doc.add_section()
new_section.page_width = Cm(29.7)
new_section.page_height = Cm(21.0)
new_section.left_margin = Inches(1.0)
new_section.right_margin = Inches(0.8)
new_section.top_margin = Inches(0.8)
new_section.bottom_margin = Inches(0.8)

add_para('14. GANTT CHART', size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=10)

add_para(
    'The Gantt chart below shows the project schedule across 4 weeks '
    'with daily-level granularity grouped by week. Overlapping colors '
    'indicate phases that run in parallel to meet the compressed timeline.',
    after=12
)

# Gantt: Task column + 4 week columns (each split into 7 days = 28 sub-cols)
# Simpler approach: Task + W1 Mon-Fri + W2 Mon-Fri + W3 Mon-Fri + W4 Mon-Fri
# Actually let's keep it simple: Task + Day1-5 per week = 21 cols total
# Even simpler: Task + 4 half-week blocks per week = Task + 8 cols
# SIMPLEST & BEST: Task + W1 + W2 + W3 + W4 with sub-tasks showing which weeks

phases = [
    # (task, weeks_active[1-4], color, is_phase_header)
    ('Phase 1: Research & Planning', [1], '2E5DA8', True),
    ('  Literature review & proposal', [1], '5B9BD5', False),
    ('  Dataset study & tool setup', [1], '5B9BD5', False),
    ('Phase 2: Data Analysis & Preprocessing', [1,2], 'C55A11', True),
    ('  Dataset download & merging', [1], 'ED7D31', False),
    ('  Exploratory Data Analysis', [1,2], 'ED7D31', False),
    ('  Data preprocessing & encoding', [2], 'F4A460', False),
    ('  Train-test split & feature scaling', [2], 'F4A460', False),
    ('Phase 3: Model Development', [2,3], '2D7D2D', True),
    ('  Feature selection', [2], '70AD47', False),
    ('  Train DT & KNN models', [2], '70AD47', False),
    ('  Train RF & XGBoost models', [2,3], '92D050', False),
    ('  Hyperparameter tuning (GridSearchCV)', [3], '92D050', False),
    ('  Model evaluation & comparison', [3], '92D050', False),
    ('Phase 4: Deployment', [3,4], 'B8860B', True),
    ('  Save model with joblib', [3], 'FFC000', False),
    ('  Build FastAPI backend', [3,4], 'FFC000', False),
    ('  Develop web form frontend', [3,4], 'FFD54F', False),
    ('  Integration & end-to-end testing', [4], 'FFD54F', False),
    ('Phase 5: Testing & Documentation', [4], '5B2C8E', True),
    ('  Functional & accuracy testing', [4], '7030A0', False),
    ('  Final report writing', [4], '9B59B6', False),
    ('  Presentation & submission', [4], '9B59B6', False),
]

num_rows = len(phases) + 1
table = doc.add_table(rows=num_rows, cols=5)
table.style = 'Table Grid'

# Header
set_cell(table.rows[0].cells[0], 'Activity / Phase', bold=True, size=10)
for w in range(1, 5):
    set_cell(table.rows[0].cells[w], f'Week {w}', bold=True, size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(table.rows[0].cells[w], 'D9E2F3')

# Set widths
table.rows[0].cells[0].width = Inches(3.5)
for w in range(1, 5):
    table.rows[0].cells[w].width = Inches(1.5)

# Data
for row_idx, (task, weeks, color, is_header) in enumerate(phases, 1):
    display_name = task.strip()
    set_cell(table.rows[row_idx].cells[0], display_name,
             bold=is_header, size=10 if is_header else 9)

    if is_header:
        shade_cell(table.rows[row_idx].cells[0], 'F2F2F2')

    for w in range(1, 5):
        if w in weeks:
            shade_cell(table.rows[row_idx].cells[w], color)

    table.rows[row_idx].cells[0].width = Inches(3.5)
    for w in range(1, 5):
        table.rows[row_idx].cells[w].width = Inches(1.5)

add_para('', after=10)

# Legend
add_para('Legend:', size=10, bold=True, before=8, after=4)

legend = doc.add_table(rows=1, cols=5)
legend.style = 'Table Grid'

legend_items = [
    ('Research & Planning', '2E5DA8'),
    ('Data Analysis', 'C55A11'),
    ('Model Development', '2D7D2D'),
    ('Deployment', 'B8860B'),
    ('Testing & Docs', '5B2C8E'),
]

for i, (label, color) in enumerate(legend_items):
    cell = legend.rows[0].cells[i]
    shade_cell(cell, color)
    set_cell(cell, label, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(255, 255, 255)

path = r'd:\Tele-Health\DevPlan_Gantt_4wk.docx'
doc.save(path)
print(f'Saved: {path}')

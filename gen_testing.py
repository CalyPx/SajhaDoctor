from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Inches(1.5)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

def add_para(text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold
    return p

def add_mixed(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    for text, bold in parts:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.bold = bold
    return p

def set_cell(cell, text, bold=False, size=10):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold

# ═══════════════════════════
# HEADING
# ═══════════════════════════
add_para('9. TESTING STRATEGY', size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=12)

# ── 9.1 Test Types ──
add_para('9.1 Test Types', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=6)

add_para('Three types of testing are planned for this project:')

add_mixed([
    ('Model Testing: ', True),
    ('This involves evaluating how well each trained model performs on data '
     'it has not seen during training. The dataset is split into 80% training '
     'and 20% testing using stratified sampling. Each model is evaluated '
     'using accuracy, macro-averaged F1-score, precision, recall, and '
     'confusion matrices. 5-fold cross-validation is used during '
     'hyperparameter tuning to make sure the results are not dependent on '
     'a single random split.', False),
])

add_mixed([
    ('Functional Testing: ', True),
    ('This covers the web application. We check whether the input form '
     'correctly accepts building details, whether the API properly processes '
     'the request, and whether the predicted damage grade is displayed back '
     'to the user. Each input field is tested with valid and invalid values '
     'to make sure the system handles edge cases without crashing.', False),
])

add_mixed([
    ('Integration Testing: ', True),
    ('This verifies that the saved model file loads correctly inside the '
     'FastAPI application and that the preprocessing steps applied during '
     'training are also applied consistently during prediction. If the model '
     'expects encoded features in a certain order, the API must send them '
     'in that exact format.', False),
])

# ── 9.2 Sample Test Cases ──
add_para('9.2 Sample Test Cases', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=6)

table = doc.add_table(rows=7, cols=5)
table.style = 'Table Grid'

headers = ['Test ID', 'Description', 'Input', 'Expected Output', 'Status']
for i, h in enumerate(headers):
    set_cell(table.rows[0].cells[i], h, bold=True, size=10)

cases = [
    ('TC-01', 'Predict damage for a low-risk building',
     '2 floors, age 5, cement foundation, RCC roof, flat land',
     'Grade 1 (Low)', '\u2014'),
    ('TC-02', 'Predict damage for a high-risk building',
     '1 floor, age 45, mud mortar stone, bamboo roof, steep slope',
     'Grade 3 (High)', '\u2014'),
    ('TC-03', 'Submit form with missing fields',
     'Empty age field',
     'Error message displayed', '\u2014'),
    ('TC-04', 'Submit form with invalid values',
     'Negative floor count',
     'Error message displayed', '\u2014'),
    ('TC-05', 'API response time',
     'Valid building input',
     'Response within 2 seconds', '\u2014'),
    ('TC-06', 'Model accuracy check',
     'Full test dataset (20%)',
     'F1-score \u2265 0.70', '\u2014'),
]

for row_idx, (tid, desc, inp, out, status) in enumerate(cases, 1):
    vals = [tid, desc, inp, out, status]
    for col_idx, val in enumerate(vals):
        set_cell(table.rows[row_idx].cells[col_idx], val, size=10)

# Set column widths
for row in table.rows:
    row.cells[0].width = Inches(0.6)
    row.cells[1].width = Inches(1.6)
    row.cells[2].width = Inches(1.8)
    row.cells[3].width = Inches(1.2)
    row.cells[4].width = Inches(0.5)

add_para('', after=4)

# ── 9.3 Validation Plan ──
add_para('9.3 Validation Plan', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=6)

add_para(
    'The primary validation metric is the macro-averaged F1-score on the '
    'held-out test set. Since the dataset has class imbalance (Grade 2 '
    'dominates), we also examine per-class precision and recall to make '
    'sure the model is not just predicting the majority class. Confusion '
    'matrices are generated for all four models to visually compare where '
    'each model makes correct and incorrect predictions. The final selected '
    'model must achieve an F1-score of at least 0.70 to be considered '
    'acceptable for deployment. If none of the models meet this threshold '
    'after tuning, additional preprocessing steps like SMOTE oversampling '
    'will be explored as a fallback.'
)

path = r'd:\Tele-Health\Testing_Strategy.docx'
doc.save(path)
print(f'Saved: {path}')

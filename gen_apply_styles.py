from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document(r'd:\Tele-Health\proposal_copy.docx')

# Define which paragraphs get Heading 1 (main sections)
heading1_indices = {
    24,   # ACKNOWLEDGEMENT
    41,   # Student's Declaration
    62,   # Abstract
    72,   # ABBREVIATIONS
    74,   # Introduction
    100,  # Problem Statement
    112,  # Objectives
    132,  # Literature Review
    154,  # 5. METHODOLOGY
    195,  # 6. REQUIREMENT DOCUMENT
    226,  # 7. SYSTEM ANALYSIS AND DESIGN
    251,  # 8. DEVELOPMENT PLAN
    273,  # 9. TESTING STRATEGY
    300,  # 10. EXPECTED RESULTS
    310,  # 11. FUTURE ENHANCEMENTS
    317,  # 12. CONCLUSION
    331,  # 13. ANNEX / APPENDIX
    357,  # 14. GANTT CHART
    368,  # References
}

# Define which paragraphs get Heading 2 (subsections)
heading2_indices = {
    155,  # 5.1 Tools and Technologies
    158,  # 5.2 Dataset
    161,  # 5.3 Model Selection Rationale
    168,  # 5.4 Training Approach
    188,  # Model Pipeline Diagram
    196,  # 6.1 Functional Requirements
    212,  # 6.2 Non-Functional Requirements
    220,  # 6.3 Technology Stack
    227,  # 7.1 System Architecture Diagram
    233,  # 7.2 Data Flow Diagram
    241,  # 7.3 Process Flowchart
    246,  # 7.4 Model Pipeline Diagram
    248,  # 7.5 Note on ER Diagram
    256,  # Responsibilities
    274,  # 9.1 Test Types
    279,  # 9.2 Sample Test Cases
    282,  # 9.3 Validation Plan
    332,  # A. Dataset Information
    334,  # B. Key Python Libraries
    336,  # C. Code Repository
    342,  # D. Sample Features
}

# Apply styles
for i, p in enumerate(doc.paragraphs):
    if i in heading1_indices:
        p.style = doc.styles['Heading 1']
        # Re-apply formatting
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
    elif i in heading2_indices:
        p.style = doc.styles['Heading 2']
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

# Now find the TOC page and insert a TOC field
# Look for "TABLE OF CONTENTS" text or insert before Introduction
# We'll insert a TOC field code that Word can update

# Find where to insert TOC — look for paragraph that says something about contents
toc_insert_idx = None
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip().upper()
    if 'TABLE OF CONTENTS' in text or 'TABLE OF CONTENT' in text:
        toc_insert_idx = i
        break

if toc_insert_idx is None:
    # Insert before ABBREVIATIONS (index 72)
    toc_insert_idx = 72

# Create TOC field
def add_toc_field(paragraph):
    """Add a TOC field to a paragraph that Word will auto-update."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    
    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-2" \\h \\z \\u '
    run2._r.append(instrText)
    
    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar2)
    
    run4 = paragraph.add_run('[Right-click here and select "Update Field" to generate Table of Contents]')
    run4.font.name = 'Times New Roman'
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(128, 128, 128)
    
    run5 = paragraph.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar3)

# Insert TOC after the TOC heading
if toc_insert_idx is not None:
    toc_para = doc.paragraphs[toc_insert_idx]
    # Add a new paragraph after the heading for the TOC field
    new_para = OxmlElement('w:p')
    toc_para._element.addnext(new_para)
    from docx.text.paragraph import Paragraph
    toc_paragraph = Paragraph(new_para, toc_para._element.getparent())
    add_toc_field(toc_paragraph)

# Save
output = r'd:\Tele-Health\Proposal_With_TOC.docx'
doc.save(output)
print(f'Saved: {output}')
print(f'Applied Heading 1 to {len(heading1_indices)} paragraphs')
print(f'Applied Heading 2 to {len(heading2_indices)} paragraphs')
print(f'TOC field inserted at paragraph {toc_insert_idx}')
print()
print('NEXT STEP: Open the file in Word, right-click the TOC area,')
print('and select "Update Field" -> "Update entire table"')

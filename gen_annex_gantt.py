from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.page_width = Cm(29.7)  # Landscape for Gantt
section.page_height = Cm(21.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(0.8)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)

def add_para(text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold
    return p

def set_cell(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold

def shade_cell(cell, color='4472C4'):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

# ═══════════════════════════════════
# SECTION 14: GANTT CHART
# ═══════════════════════════════════

add_para('14. GANTT CHART', size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=10)

add_para(
    'The following Gantt chart illustrates the project timeline across '
    '12 weeks. Each shaded cell represents a week during which that '
    'particular activity is scheduled. Both team members collaborate '
    'on research and documentation, while data work and deployment are '
    'handled by Rohit Poudel, and model training and evaluation are '
    'handled by Keshav KC.',
    after=12
)

# Gantt table: 1 header row + rows for tasks, columns = Task + W1-W12
phases = [
    ('Phase 1: Research & Planning', [1,2], '4472C4'),
    ('  Literature review', [1,2], '5B9BD5'),
    ('  Proposal writing', [1,2], '5B9BD5'),
    ('  Tool & environment setup', [2], '5B9BD5'),
    ('Phase 2: Data Analysis', [3,4], 'ED7D31'),
    ('  Dataset download & merging', [3], 'F4A460'),
    ('  Exploratory Data Analysis', [3,4], 'F4A460'),
    ('  Data preprocessing & encoding', [4], 'F4A460'),
    ('Phase 3: Model Development', [5,6,7], '70AD47'),
    ('  Feature selection', [5], '92D050'),
    ('  Model training (DT, KNN, RF, XGBoost)', [5,6], '92D050'),
    ('  Hyperparameter tuning', [6,7], '92D050'),
    ('  Model evaluation & comparison', [7], '92D050'),
    ('Phase 4: Deployment', [8,9], 'FFC000'),
    ('  Save model with joblib', [8], 'FFD54F'),
    ('  Build FastAPI backend', [8,9], 'FFD54F'),
    ('  Develop web frontend', [9], 'FFD54F'),
    ('Phase 5: Testing & Documentation', [10,11,12], '7030A0'),
    ('  Functional & integration testing', [10,11], '9B59B6'),
    ('  Final report writing', [10,11,12], '9B59B6'),
    ('  Presentation preparation', [12], '9B59B6'),
    ('  Project submission', [12], '9B59B6'),
]

num_rows = len(phases) + 1
num_cols = 13  # Task + W1-W12

table = doc.add_table(rows=num_rows, cols=num_cols)
table.style = 'Table Grid'

# Header row
set_cell(table.rows[0].cells[0], 'Activity / Phase', bold=True, size=9)
for w in range(1, 13):
    set_cell(table.rows[0].cells[w], f'W{w}', bold=True, size=8,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(table.rows[0].cells[w], 'D9E2F3')

# Set task column width
table.rows[0].cells[0].width = Inches(2.8)
for w in range(1, 13):
    table.rows[0].cells[w].width = Inches(0.55)

# Data rows
for row_idx, (task, weeks, color) in enumerate(phases, 1):
    is_phase = not task.startswith('  ')
    set_cell(table.rows[row_idx].cells[0], task.strip(),
             bold=is_phase, size=9 if is_phase else 8)

    if is_phase:
        shade_cell(table.rows[row_idx].cells[0], 'F2F2F2')

    for w in range(1, 13):
        if w in weeks:
            shade_cell(table.rows[row_idx].cells[w], color)
        # Keep cell empty for non-shaded weeks

    table.rows[row_idx].cells[0].width = Inches(2.8)
    for w in range(1, 13):
        table.rows[row_idx].cells[w].width = Inches(0.55)

add_para('', after=8)

# Legend
add_para('Legend:', size=10, bold=True, before=8, after=4)

legend_table = doc.add_table(rows=1, cols=5)
legend_table.style = 'Table Grid'

legend_items = [
    ('Research & Planning', '4472C4'),
    ('Data Analysis', 'ED7D31'),
    ('Model Development', '70AD47'),
    ('Deployment', 'FFC000'),
    ('Testing & Docs', '7030A0'),
]

for i, (label, color) in enumerate(legend_items):
    cell = legend_table.rows[0].cells[i]
    shade_cell(cell, color)
    set_cell(cell, label, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    # Make text white for dark backgrounds
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(255, 255, 255)


# ═══════════════════════════════════
# NEW PAGE — SECTION 13: ANNEX
# ═══════════════════════════════════

doc.add_page_break()

# Switch back to portrait for Annex
new_section = doc.add_section()
new_section.page_width = Cm(21.0)
new_section.page_height = Cm(29.7)
new_section.left_margin = Inches(1.5)
new_section.right_margin = Inches(1.0)
new_section.top_margin = Inches(1.0)
new_section.bottom_margin = Inches(1.0)

add_para('13. ANNEX / APPENDIX', size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=12)

# A. Dataset
add_para('A. Dataset Information', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=6)

ds_table = doc.add_table(rows=6, cols=2)
ds_table.style = 'Table Grid'

ds_data = [
    ('Dataset Name', 'Richter\u2019s Predictor: Modeling Earthquake Damage'),
    ('Source', 'DrivenData (https://www.drivendata.org/competitions/57/nepal-earthquake/)'),
    ('Collected By', 'Central Bureau of Statistics, Nepal & Kathmandu Living Labs'),
    ('Total Records', '260,601 buildings'),
    ('Features', '39 (structural, geographic, categorical, and binary)'),
    ('Target Variable', 'damage_grade (Grade 1, Grade 2, Grade 3)'),
]

for row_idx, (label, value) in enumerate(ds_data):
    set_cell(ds_table.rows[row_idx].cells[0], label, bold=True, size=11)
    set_cell(ds_table.rows[row_idx].cells[1], value, size=11)

ds_table.rows[0].cells[0].width = Inches(1.5)
ds_table.rows[0].cells[1].width = Inches(4.0)

add_para('', after=6)

# B. Key Libraries
add_para('B. Key Python Libraries and Versions', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=10, after=6)

lib_table = doc.add_table(rows=9, cols=3)
lib_table.style = 'Table Grid'

lib_headers = ['Library', 'Version', 'Purpose']
for i, h in enumerate(lib_headers):
    set_cell(lib_table.rows[0].cells[i], h, bold=True, size=11)

libs = [
    ('Python', '3.11+', 'Programming language'),
    ('scikit-learn', '1.3+', 'ML model training, evaluation, GridSearchCV'),
    ('xgboost', '2.0+', 'XGBoost classifier implementation'),
    ('pandas', '2.0+', 'Data loading and manipulation'),
    ('numpy', '1.24+', 'Numerical computation'),
    ('matplotlib', '3.7+', 'Data visualization and charts'),
    ('seaborn', '0.12+', 'Statistical plots and heatmaps'),
    ('fastapi', '0.100+', 'REST API backend for web deployment'),
]

for row_idx, (lib, ver, purpose) in enumerate(libs, 1):
    set_cell(lib_table.rows[row_idx].cells[0], lib, size=11)
    set_cell(lib_table.rows[row_idx].cells[1], ver, size=11)
    set_cell(lib_table.rows[row_idx].cells[2], purpose, size=11)

add_para('', after=6)

# C. GitHub
add_para('C. Code Repository', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=10, after=6)

add_para(
    'All source code for this project will be hosted on a public GitHub '
    'repository. The repository will include a README.md file with setup '
    'instructions, a requirements.txt listing all dependencies, the '
    'Jupyter Notebook used for training and evaluation, and the FastAPI '
    'application code for the web deployment.'
)

add_para('GitHub URL: [To be added after repository creation]',
         bold=True, after=6)

add_para('', after=6)

# D. Sample feature list
add_para('D. Sample Features from Dataset', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=10, after=6)

feat_table = doc.add_table(rows=11, cols=3)
feat_table.style = 'Table Grid'

feat_headers = ['Feature Name', 'Type', 'Description']
for i, h in enumerate(feat_headers):
    set_cell(feat_table.rows[0].cells[i], h, bold=True, size=10)

feats = [
    ('geo_level_1_id', 'Integer', 'Geographic region identifier (level 1)'),
    ('count_floors_pre_eq', 'Integer', 'Number of floors before the earthquake'),
    ('age', 'Integer', 'Age of the building in years'),
    ('area_percentage', 'Integer', 'Normalized area of the building footprint'),
    ('height_percentage', 'Integer', 'Normalized height of the building'),
    ('foundation_type', 'Categorical', 'Type of foundation (h, i, r, u, w)'),
    ('roof_type', 'Categorical', 'Type of roof (n, q, x)'),
    ('ground_floor_type', 'Categorical', 'Type of ground floor (f, m, v, x, z)'),
    ('land_surface_condition', 'Categorical', 'Surface condition (n, o, t)'),
    ('has_superstructure_mud_mortar_stone', 'Binary', 'Whether walls use mud mortar stone'),
]

for row_idx, (name, ftype, desc) in enumerate(feats, 1):
    set_cell(feat_table.rows[row_idx].cells[0], name, size=9)
    set_cell(feat_table.rows[row_idx].cells[1], ftype, size=10)
    set_cell(feat_table.rows[row_idx].cells[2], desc, size=10)

path = r'd:\Tele-Health\Annex_Gantt.docx'
doc.save(path)
print(f'Saved: {path}')

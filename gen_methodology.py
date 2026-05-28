from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Inches(1.5)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

def add_para(text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=8, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    return p

def add_mixed(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    for text, bold, italic in parts:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.bold = bold
        r.font.italic = italic
    return p

# ═══════════════════════════════════════
# HEADING
# ═══════════════════════════════════════
add_para('5. METHODOLOGY', size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=12)

# ── 5.1 Tools ──
add_para('5.1 Tools and Technologies', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=10, after=8)

add_para(
    'This project is developed entirely in Python 3.x using open-source '
    'libraries. The core machine learning workflow uses scikit-learn for '
    'model training, evaluation, and hyperparameter tuning, along with '
    'XGBoost for the gradient boosting implementation. Data manipulation '
    'and preprocessing are handled using pandas and NumPy. Exploratory data '
    'analysis and visualization use matplotlib and seaborn. For the '
    'deployment component, FastAPI is used to build a lightweight REST API '
    'that serves predictions through a web interface. The frontend is built '
    'using HTML, CSS, and JavaScript. Jupyter Notebook is used for the '
    'experimentation and model development phase, while the final trained '
    'model is serialized using joblib for integration with the web application.'
)

# ── 5.2 Dataset ──
add_para('5.2 Dataset', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=10, after=8)

add_para(
    'The dataset used in this project is sourced from the '
    '\u201cRichter\u2019s Predictor: Modeling Earthquake Damage\u201d '
    'competition hosted by DrivenData '
    '(https://www.drivendata.org/competitions/57/nepal-earthquake/). '
    'The data was originally collected by the Central Bureau of Statistics '
    'of Nepal and Kathmandu Living Labs through household surveys conducted '
    'after the 7.8 Mw Gorkha earthquake on April 25, 2015. The dataset '
    'contains 260,601 building records with 39 features describing structural '
    'properties such as number of floors, age of building, floor area, '
    'height, foundation type, roof type, ground floor type, land surface '
    'condition, and construction material. The target variable is '
    'damage_grade, which classifies each building into one of three ordinal '
    'levels: Grade 1 (low damage), Grade 2 (medium damage), and Grade 3 '
    '(high damage). The dataset is publicly available and does not require '
    'any ethical clearance for academic use.'
)

# ── 5.3 Model Selection ──
add_para('5.3 Model Selection Rationale', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=10, after=8)

add_para(
    'Four supervised classification algorithms are selected for comparative '
    'evaluation:'
)

models = [
    ('Decision Tree Classifier',
     ' \u2014 selected as a baseline model due to its interpretability and '
     'ability to handle both categorical and numerical features without scaling.'),
    ('K-Nearest Neighbor (KNN)',
     ' \u2014 selected as a distance-based, non-parametric classifier to '
     'evaluate how spatial similarity among building features contributes '
     'to damage prediction.'),
    ('Random Forest Classifier',
     ' \u2014 selected because the literature consistently identifies it as '
     'the best-performing model for this type of task (Ghimire et al., 2022; '
     'Bhatta & Dang, 2023; Onur et al., 2022).'),
    ('XGBoost Classifier',
     ' \u2014 selected as an advanced gradient boosting method that includes '
     'built-in regularization to reduce overfitting, and has shown strong '
     'performance on structured tabular data.'),
]

for i, (name, desc) in enumerate(models, 1):
    add_mixed([
        (f'{i}. ', False, False),
        (name, True, False),
        (desc, False, False),
    ])

add_para(
    'These four models represent a progression from simple to complex, '
    'allowing a meaningful comparison of accuracy, training time, and '
    'generalization ability.'
)

# ── 5.4 Training Approach ──
add_para('5.4 Training Approach', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=10, after=8)

add_para(
    'The end-to-end machine learning pipeline follows these steps:'
)

steps = [
    ('Step 1 \u2014 Data Collection: ',
     'Download the dataset from DrivenData and merge the feature and target files.'),
    ('Step 2 \u2014 Exploratory Data Analysis: ',
     'Analyze feature distributions, class balance, correlations, and missing '
     'values. Visualize damage grade distribution across building types and '
     'geographic regions.'),
    ('Step 3 \u2014 Data Preprocessing: ',
     'Encode categorical variables using one-hot encoding or label encoding. '
     'Handle missing values where applicable. Apply feature scaling for KNN. '
     'Address class imbalance using stratified train-test splitting with an '
     '80:20 ratio.'),
    ('Step 4 \u2014 Feature Selection: ',
     'Evaluate feature importance using Random Forest-based importance scores. '
     'Remove low-contribution features to reduce noise and improve model efficiency.'),
    ('Step 5 \u2014 Model Training: ',
     'Train all four models on the training set. Perform hyperparameter tuning '
     'using GridSearchCV with 5-fold cross-validation.'),
    ('Step 6 \u2014 Model Evaluation: ',
     'Compare models using micro-averaged F1-score (primary metric), accuracy, '
     'precision, recall, and confusion matrices. Select the best-performing model.'),
    ('Step 7 \u2014 Deployment: ',
     'Serialize the best model using joblib. Build a FastAPI backend that '
     'accepts building features as input and returns the predicted damage '
     'grade. Connect the API to a simple web interface for user interaction.'),
]

for label, desc in steps:
    add_mixed([
        (label, True, False),
        (desc, False, False),
    ])

# ── 5.5 Pipeline Diagram ──
add_para('5.5 Model Pipeline Diagram', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=10, after=8)

# Insert the pipeline image
import os
img_path = r'C:\Users\Gaming F16\.gemini\antigravity\brain\4ee70f2d-ccac-4272-98d6-3026c937c32e\ml_pipeline_diagram_1776918412604.png'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run()
r.add_picture(img_path, width=Inches(5.0))

# Caption
add_para('Figure 1: End-to-End Machine Learning Pipeline for Earthquake '
         'Building Damage Prediction', size=10, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=8)

path = r'd:\Tele-Health\Methodology_Section.docx'
doc.save(path)
print(f'Saved: {path}')

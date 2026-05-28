from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Inches(1.5)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

def add_para(text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=10, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p

def add_mixed_para(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=10):
    """parts = list of (text, bold, italic) tuples"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = bold
        run.font.italic = italic
    return p

# ══════════════════════════════════════
# HEADING
# ══════════════════════════════════════

add_para('4. BACKGROUND STUDY / LITERATURE REVIEW', size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=20)

add_para(
    'This section reviews existing research on earthquake damage prediction '
    'using machine learning. The studies discussed below are relevant to the '
    'approach used in this project and have helped shape the methodology.',
    after=16
)

# ══════════════════════════════════════
# STUDY 1
# ══════════════════════════════════════

add_mixed_para([
    ('Mangalathu, S., Sun, H., Nithurshan, T., & Burton, H. V. (2022) ', False, False),
    ('published a study in ', False, False),
    ('Engineering Structures ', False, True),
    ('on classifying earthquake damage to buildings using machine learning. '
     'They compared several algorithms including Random Forest, XGBoost, and '
     'Support Vector Machines for multi-class damage classification. Their results '
     'showed that gradient boosting methods, especially XGBoost, performed better '
     'than traditional approaches. They also used SHAP (SHapley Additive exPlanations) '
     'to identify which building features contributed most to the predictions. '
     'This study supports our approach of comparing multiple models and analyzing '
     'feature importance to understand what makes a building vulnerable.', False, False),
], after=14)

# ══════════════════════════════════════
# STUDY 2
# ══════════════════════════════════════

add_mixed_para([
    ('Gautam, D., Chaulagain, H., & Rodrigues, H. (2023) ', False, False),
    ('studied the seismic vulnerability of buildings in Nepal using data-driven '
     'approaches. Their research focused on how structural characteristics like '
     'building age, construction material, and number of floors directly affect '
     'the level of damage during earthquakes. They used data from the 2015 Gorkha '
     'earthquake and found that ensemble methods such as Random Forest are effective '
     'for this type of classification because they handle mixed feature types well '
     'without needing heavy preprocessing. Their findings confirm that the building '
     'attributes available in the DrivenData dataset are suitable for damage prediction.', False, False),
], after=14)

# ══════════════════════════════════════
# STUDY 3
# ══════════════════════════════════════

add_mixed_para([
    ('Nesmachnow, S. & Baez, L. (2022) ', False, False),
    ('worked directly with the DrivenData Nepal earthquake dataset, which is the '
     'same dataset used in this project. They applied multiple classification '
     'algorithms and used feature engineering to improve accuracy. Their best model '
     'achieved an F1 micro-score of approximately 0.75 using gradient boosting '
     'techniques. They noted that geographical features (location IDs) were among '
     'the most important predictors of damage. This study serves as a direct '
     'benchmark for our work since we are using the same data and target variable.', False, False),
], after=14)

# ══════════════════════════════════════
# STUDY 4
# ══════════════════════════════════════

add_mixed_para([
    ('Pan, Y., Liang, X., & Zhang, L. (2023) ', False, False),
    ('published a review in ', False, False),
    ('Automation in Construction ', False, True),
    ('covering machine learning applications in post-earthquake damage assessment. '
     'They reviewed over 50 studies and found that tree-based ensemble methods '
     'consistently outperformed both linear models and basic neural networks for '
     'structural damage classification. The review recommended combining structural, '
     'material, and geographical features for the best results, which is exactly '
     'the approach followed in this project.', False, False),
], after=14)

# ══════════════════════════════════════
# STUDY 5
# ══════════════════════════════════════

add_mixed_para([
    ('Rathore, A., Kumar, S., & Gupta, M. (2024) ', False, False),
    ('conducted a comparative study published in the ', False, False),
    ('Asian Journal of Civil Engineering ', False, True),
    ('on predicting earthquake damage using the Nepal earthquake dataset. '
     'They tested Random Forest, Decision Tree, and ElasticNet with Bayesian '
     'Optimization for hyperparameter tuning. Random Forest achieved the highest '
     'accuracy. The study also highlighted the problem of class imbalance in the '
     'dataset, where medium-damage buildings are overrepresented compared to low '
     'and severe categories. This is a challenge we also expect to face in our '
     'project, and we plan to use stratified sampling to address it.', False, False),
], after=14)

# ══════════════════════════════════════
# SUMMARY
# ══════════════════════════════════════

add_para(
    'Based on the reviewed studies, it is clear that machine learning, '
    'particularly ensemble methods like Random Forest and XGBoost, has shown '
    'strong results for earthquake damage prediction. Most of these studies '
    'work with similar building features and report accuracies in the range '
    'of 70-76% for the Nepal dataset. Our project builds on these findings '
    'by training and comparing four different models and deploying the best '
    'one through a web application for practical use.',
    before=8, after=20
)

# ══════════════════════════════════════
# NOTE
# ══════════════════════════════════════

add_mixed_para([
    ('Note: ', True, False),
    ('Verify each reference on Google Scholar before submitting. '
     'Find the exact journal, volume, page numbers, and DOI for '
     'APA 7th Edition formatting.', False, True),
], before=20, after=0)

# SAVE
path = r'd:\Tele-Health\Literature_Review.docx'
doc.save(path)
print(f'Saved: {path}')

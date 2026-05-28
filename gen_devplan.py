from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Inches(1.5)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

def add_para(text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold
    return p

# Heading
add_para('8. DEVELOPMENT PLAN', size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=12)

add_para(
    'The project is planned over a 10-week period divided into five phases. '
    'Since this is an individual project, all tasks are carried out by the '
    'student under the guidance of the project supervisor. The table below '
    'outlines the phases, their key activities, expected milestones, and timeline.'
)

# Table
table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'

headers = ['Phase', 'Activities', 'Milestone', 'Timeline']
data = [
    ('Phase 1:\nResearch &\nPlanning',
     'Literature review, proposal writing, dataset study, tool setup',
     'Proposal approved, development environment ready',
     'Week 1\u20132'),

    ('Phase 2:\nData Analysis &\nPreprocessing',
     'Download dataset, perform EDA, handle missing values, '
     'encode features, split data',
     'Clean and processed dataset ready for training',
     'Week 3\u20134'),

    ('Phase 3:\nModel\nDevelopment',
     'Train Decision Tree, KNN, Random Forest, and XGBoost models. '
     'Perform hyperparameter tuning using GridSearchCV. '
     'Run feature importance analysis',
     'All four models trained and evaluated with comparison metrics',
     'Week 5\u20137'),

    ('Phase 4:\nDeployment',
     'Save best model with joblib, build FastAPI backend, '
     'develop web form frontend, integrate prediction pipeline',
     'Working web application that accepts input and returns damage grade',
     'Week 8\u20139'),

    ('Phase 5:\nTesting &\nDocumentation',
     'Test the web app with sample inputs, verify prediction accuracy, '
     'write final report, prepare presentation',
     'Final report submitted, project demonstrated',
     'Week 10'),
]

# Header row
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    r = cell.paragraphs[0].add_run(h)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    r.font.bold = True

# Data rows
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        cell = table.rows[row_idx].cells[col_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if col_idx == 0:
            r = p.add_run(val)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
            r.font.bold = True
        else:
            r = p.add_run(val)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

# Set column widths
for row in table.rows:
    row.cells[0].width = Inches(1.2)
    row.cells[1].width = Inches(2.3)
    row.cells[2].width = Inches(2.0)
    row.cells[3].width = Inches(0.8)

add_para('', after=6)

# Responsibilities
add_para('Responsibilities', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=6)

add_para(
    'Since this is an individual microdegree project, all phases are handled '
    'by the student. The project supervisor provides guidance during the '
    'research, design review, and final evaluation stages.'
)

path = r'd:\Tele-Health\Development_Plan.docx'
doc.save(path)
print(f'Saved: {path}')

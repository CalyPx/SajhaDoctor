from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Inches(1.5)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

def add_mixed(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=10):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    for text, bold, italic in parts:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.bold = bold
        r.font.italic = italic
    return p

def add_para(text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=10, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    return p

# ══════════════════════════════
# HEADING
# ══════════════════════════════

add_para('4. BACKGROUND STUDY / LITERATURE REVIEW', size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=18)

# ── Intro paragraph ──
add_para(
    'Earthquake damage prediction has become an active area of research '
    'in recent years, especially after large-scale post-earthquake datasets '
    'became publicly available. Traditional methods of assessing building '
    'vulnerability rely on manual inspection by engineers, which is slow, '
    'expensive, and hard to scale across entire regions. Several researchers '
    'have explored machine learning as a faster and more practical alternative. '
    'This section reviews key studies that are directly related to the approach '
    'taken in this project.',
    after=14
)

# ── Study 1: Ghimire et al. (2022) ──
add_mixed([
    ('Ghimire et al. (2022) analyzed building damage data from the 2015 Gorkha '
     'earthquake at a regional scale by combining structural parameters from the '
     'Nepal Building Damage Portfolio with ground motion data from the USGS ShakeMap '
     'tool. They tested multiple machine learning models and found that Random Forest '
     'produced the best results when ground motion intensity was included as an '
     'additional feature. Their study achieved an accuracy of around 68% for multi-class '
     'damage prediction and showed that even a small number of key features like building '
     'age, height, and construction type can provide useful predictions. This is relevant '
     'to our project because we also rely on similar building attributes from the same '
     'earthquake event.', False, False),
], after=14)

# ── Study 2: Bhatta & Dang (2023) ──
add_mixed([
    ('Bhatta and Dang (2023) focused specifically on reinforced concrete buildings '
     'damaged during the 2015 Nepal earthquake. They evaluated K-Nearest Neighbor, '
     'Random Forest, Decision Tree, Support Vector Machine, and Artificial Neural '
     'Network models using actual post-earthquake field survey data. Their study '
     'validated that simulation-based model training can produce results close to '
     'real-world damage patterns. The selection of algorithms in their study closely '
     'matches the ones we use in this project, which gives us confidence that '
     'Decision Tree, KNN, Random Forest, and XGBoost are appropriate choices for '
     'this type of classification task.', False, False),
], after=14)

# ── Study 3: Adi et al. (2020) ──
add_mixed([
    ('Adi et al. (2020) worked directly with the DrivenData Nepal earthquake '
     'dataset, which is the same dataset used in this project. They applied '
     'Random Forest and Gradient Boosting Classifier with different preprocessing '
     'steps and found that even simple improvements like outlier handling and '
     'hyperparameter tuning can noticeably improve results. With parameter tuning, '
     'their Random Forest model achieved an F1-score of 72.95%, which improved '
     'to 74.42% using Gradient Boosting with tuned hyperparameters. These results '
     'serve as a direct benchmark for our project since we are working with the '
     'same data and target variable.', False, False),
], after=14)

# ── Study 4: Poudyal & Adhikari (2025) ──
add_mixed([
    ('A more recent study by Poudyal and Adhikari (2025) also used the 2015 '
     'Gorkha earthquake dataset and tested Random Forest with oversampling '
     'techniques like SMOTE, along with XGBoost, Support Vector Machine, and '
     'Multilayer Perceptron. Their best model, Random Forest with SMOTE, achieved '
     'a macro F1-score of 0.67. Their study highlighted the problem of class '
     'imbalance in the dataset, where medium-damage buildings are overrepresented '
     'compared to the low and severe categories. This is a challenge we also '
     'expect to face, and we plan to address it through stratified sampling '
     'during train-test splitting.', False, False),
], after=14)

# ── Study 5: Celebi et al. (2025) ──
add_mixed([
    ('Looking beyond Nepal, Celebi et al. (2025) evaluated nine machine learning '
     'models on a much larger dataset of 965,270 buildings damaged during the '
     'February 2023 earthquake in Turkey. Random Forest achieved the highest '
     'classification accuracy of 93% in their study, while KNN showed competitive '
     'accuracy but struggled with predicting collapse cases specifically. This '
     'result from a different seismic context confirms that ensemble methods like '
     'Random Forest and XGBoost generalize well for multi-class building damage '
     'prediction, not just for the Nepal dataset.', False, False),
], after=14)

# ── Study 6: Kartal et al. (2025) ──
add_mixed([
    ('Kartal et al. (2025) focused on gradient boosting classifiers for '
     'reinforced concrete structures damaged in the Nepal earthquake. They tested '
     'Gradient Boosting Classifier, XGBoost, and LightGBM, and found that '
     'building-type-specific modeling combined with class balancing techniques '
     'improved prediction reliability. Their work further supports our decision '
     'to include XGBoost as one of the candidate models in this project.', False, False),
], after=14)

# ── Summary paragraph ──
add_para(
    'The reviewed studies show a clear pattern: ensemble methods like Random '
    'Forest and XGBoost consistently outperform simpler models for earthquake '
    'damage classification. Reported F1-scores on the Nepal dataset typically '
    'fall in the range of 0.67 to 0.75, depending on preprocessing and tuning. '
    'However, most of these studies treat the prediction task as a standalone '
    'research exercise and do not produce a tool that non-technical users can '
    'actually use. This project aims to fill that gap by not only comparing '
    'multiple models but also deploying the best one through a web application '
    'that can be used for practical earthquake preparedness planning.',
    after=20
)

# ══════════════════════════════
# REFERENCES
# ══════════════════════════════

add_para('References', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
         before=20, after=14)

refs = [
    'Adi, S. P., Adishesha, V. B., Bharadwaj, K. V., & Narayan, A. (2020). '
    'Earthquake damage prediction using Random Forest and Gradient Boosting Classifier. '
    'American Journal of Biological and Environmental Statistics, 6(3), 58-63. '
    'https://doi.org/10.11648/j.ajbes.20200603.14',

    'Bhatta, S., & Dang, J. (2023). Seismic damage prediction of RC buildings '
    'using machine learning. Earthquake Engineering & Structural Dynamics, 52(9), '
    '2692-2711. https://doi.org/10.1002/eqe.3907',

    'Celebi, O., Arslan, M. H., & Ceylan, M. (2025). A big data-enabled decision '
    'support model for post-earthquake damage classification of RC buildings. '
    'Journal of Earthquake Engineering, 29(10), 2192-2214. '
    'https://doi.org/10.1080/13632469.2025.2505974',

    'Ghimire, S., Gueguen, P., Giffard-Roisin, S., & Schorlemmer, D. (2022). '
    'Testing machine learning models for seismic damage prediction at a regional '
    'scale using the 2015 Gorkha Nepal earthquake. Earthquake Spectra, 38(4), '
    '2970-2993. https://doi.org/10.1177/87552930221106495',

    'Kartal, S., Kocaman, I., & Ozkilic, Y. O. (2025). Seismic damage assessment '
    'of RC structures after the 2015 Gorkha, Nepal earthquake using gradient '
    'boosting classifiers. Buildings, 15(19), 3577. '
    'https://doi.org/10.3390/buildings15193577',

    'Poudyal, R., & Adhikari, P. (2025). Enhancing earthquake preparedness in '
    'Nepal through machine learning-based damage prediction models. ResearchGate '
    'Preprint. https://doi.org/10.13140/RG.2.2.xxxxx',
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    r = p.add_run(ref)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)

# SAVE
path = r'd:\Tele-Health\Literature_Review_Final.docx'
doc.save(path)
print(f'Saved: {path}')
